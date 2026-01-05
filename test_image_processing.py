#!/usr/bin/env python3
"""
Image Processing Test
Tests the image handling in Word documents
"""

import base64
import io
from docx import Document
from docx.shared import Inches

def test_image_processing():
    """Test image processing functionality"""
    print("🖼️  Testing Image Processing...")
    
    # Create a simple 1x1 pixel PNG image (base64)
    test_image_b64 = "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAYAAAAfFcSJAAAADUlEQVR42mNkYPhfDwAChwGA60e6kgAAAABJRU5ErkJggg=="
    
    try:
        # Test 1: Base64 decoding
        print("  📝 Testing base64 decoding...")
        image_bytes = base64.b64decode(test_image_b64)
        print(f"  ✅ Decoded {len(image_bytes)} bytes")
        
        # Test 2: BytesIO stream creation
        print("  📝 Testing BytesIO stream...")
        image_stream = io.BytesIO(image_bytes)
        print(f"  ✅ Created stream at position {image_stream.tell()}")
        
        # Test 3: Stream reset
        print("  📝 Testing stream reset...")
        image_stream.seek(0)
        print(f"  ✅ Reset stream to position {image_stream.tell()}")
        
        # Test 4: Word document image insertion
        print("  📝 Testing Word document insertion...")
        doc = Document()
        para = doc.add_paragraph()
        run = para.add_run()
        
        # Reset stream before adding to document
        image_stream.seek(0)
        run.add_picture(image_stream, width=Inches(1))
        print("  ✅ Successfully added image to Word document")
        
        # Test 5: Save document
        print("  📝 Testing document save...")
        test_file = "test_image_document.docx"
        doc.save(test_file)
        print(f"  ✅ Saved test document: {test_file}")
        
        print("  🎉 All image processing tests PASSED")
        return True
        
    except Exception as e:
        print(f"  ❌ Image processing test failed: {str(e)}")
        return False

def test_data_url_processing():
    """Test data URL processing"""
    print("\n🔗 Testing Data URL Processing...")
    
    try:
        # Test data URL with prefix
        data_url = "data:image/png;base64,iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAYAAAAfFcSJAAAADUlEQVR42mNkYPhfDwAChwGA60e6kgAAAABJRU5ErkJggg=="
        
        print("  📝 Testing data URL prefix removal...")
        if data_url.startswith('data:image'):
            base64_string = data_url.split(',')[1]
            print("  ✅ Successfully removed data URL prefix")
        else:
            base64_string = data_url
            print("  ✅ No prefix to remove")
        
        # Test decoding
        print("  📝 Testing base64 decoding...")
        image_bytes = base64.b64decode(base64_string)
        print(f"  ✅ Decoded {len(image_bytes)} bytes")
        
        print("  🎉 Data URL processing tests PASSED")
        return True
        
    except Exception as e:
        print(f"  ❌ Data URL processing test failed: {str(e)}")
        return False

def main():
    """Run image processing tests"""
    print("🖼️  IMAGE PROCESSING TEST SUITE")
    print("=" * 40)
    
    test1 = test_image_processing()
    test2 = test_data_url_processing()
    
    print("\n" + "=" * 40)
    print("📊 IMAGE TEST RESULTS")
    print("=" * 40)
    
    if test1 and test2:
        print("🎉 ALL IMAGE TESTS PASSED!")
        print("✅ Images should now work in Word documents")
        return True
    else:
        print("❌ Some image tests failed")
        print("⚠️  Image functionality may not work properly")
        return False

if __name__ == "__main__":
    success = main()
    exit(0 if success else 1)