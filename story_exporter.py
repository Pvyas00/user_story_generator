import os
import tempfile
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.shared import OxmlElement, qn
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib import colors
from reportlab.platypus import Image as RLImage
from reportlab.graphics.shapes import Drawing, Circle, Rect, String
from reportlab.graphics.charts.piecharts import Pie
from reportlab.graphics import renderPDF
try:
    from PIL import Image, ImageDraw, ImageFont
except ImportError:
    Image = ImageDraw = ImageFont = None
import io

class StoryExporter:
    def __init__(self):
        self.temp_dir = tempfile.gettempdir()
        self.logo_path = os.path.join('static', 'images', 'anand_rathi_logo.png')
        
        # Corporate Color Scheme
        self.colors = {
            'primary_blue': RGBColor(30, 64, 175),      # #1e40af
            'success_green': RGBColor(5, 150, 105),     # #059669
            'warning_orange': RGBColor(217, 119, 6),    # #d97706
            'error_red': RGBColor(220, 38, 38),         # #dc2626
            'light_gray': RGBColor(248, 250, 252),      # #f8fafc
            'medium_gray': RGBColor(100, 116, 139),     # #64748b
            'dark_gray': RGBColor(30, 41, 59),          # #1e293b
            'accent_purple': RGBColor(147, 51, 234)     # #9333ea
        }
        
        # Smart Templates
        self.templates = {
            'executive': ['business_goal', 'actor', 'risks', 'dependencies'],
            'technical': ['functional_flow', 'security', 'validations', 'dependencies'],
            'agile': ['business_goal', 'actor', 'acceptance_criteria', 'functional_flow'],
            'compliance': ['security', 'risks', 'validations', 'dependencies']
        }
        
        # Section Colors
        self.section_colors = {
            'business_goal': self.colors['primary_blue'],
            'actor': self.colors['success_green'],
            'trigger': self.colors['warning_orange'],
            'preconditions': self.colors['medium_gray'],
            'functional_flow': self.colors['primary_blue'],
            'validations': self.colors['success_green'],
            'acceptance_criteria': self.colors['primary_blue'],
            'security': self.colors['error_red'],
            'dependencies': self.colors['warning_orange'],
            'risks': self.colors['accent_purple']
        }
    
    def export_story(self, story_data, format_type, template='agile', coverage_data=None):
        """Export story to specified format with template"""
        if format_type == 'word':
            return self._export_word(story_data, template, coverage_data)
        elif format_type == 'pdf':
            return self._export_pdf(story_data, template, coverage_data)
        elif format_type == 'png':
            return self._export_png(story_data)
        else:
            raise ValueError(f"Unsupported format: {format_type}")
    
    def _export_word(self, story_data, template, coverage_data):
        """Export story to Word document with smart template"""
        doc = Document()
        
        # Cover Page
        self._add_word_cover_page(doc, story_data)
        doc.add_page_break()
        
        # Executive Dashboard (if coverage data available)
        if coverage_data:
            self._add_word_dashboard(doc, coverage_data)
            doc.add_page_break()
        
        # Executive Summary
        self._add_word_executive_summary(doc, story_data, coverage_data)
        doc.add_page_break()
        
        # Main Content based on template
        self._add_word_content(doc, story_data, template)
        
        # Save document
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"enterprise_user_story_{timestamp}.docx"
        filepath = os.path.join(self.temp_dir, filename)
        doc.save(filepath)
        
        return filepath
    
    def _add_word_cover_page(self, doc, story_data):
        """Add professional cover page with corporate styling"""
        # Logo with border
        if os.path.exists(self.logo_path):
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para.add_run()
            run.add_picture(self.logo_path, width=Inches(3.0))
        
        doc.add_paragraph()
        
        # Title with professional styling
        title = doc.add_heading('Enterprise User Story', level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title.runs[0]
        title_run.font.size = Pt(32)
        title_run.font.color.rgb = self.colors['primary_blue']
        title_run.font.name = 'Calibri'
        title_run.bold = True
        
        # Business Goal with styling
        if story_data.get('business_goal'):
            subtitle = doc.add_paragraph(story_data['business_goal'][:120])
            subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
            subtitle_run = subtitle.runs[0]
            subtitle_run.font.size = Pt(16)
            subtitle_run.font.color.rgb = self.colors['dark_gray']
            subtitle_run.font.name = 'Calibri'
            subtitle_run.italic = True
        
        doc.add_paragraph()
        doc.add_paragraph()
        
        # Professional metadata table
        meta_table = doc.add_table(rows=4, cols=2)
        meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Table content with styling
        cells = [
            ('Created Date:', datetime.now().strftime("%B %d, %Y")),
            ('Document Type:', 'Enterprise User Story'),
            ('Version:', '1.0'),
            ('Status:', 'Draft')
        ]
        
        for i, (label, value) in enumerate(cells):
            row = meta_table.rows[i]
            row.cells[0].text = label
            row.cells[1].text = value
            
            # Style header cells
            label_run = row.cells[0].paragraphs[0].runs[0]
            label_run.font.bold = True
            label_run.font.color.rgb = self.colors['primary_blue']
            
            # Style value cells
            value_run = row.cells[1].paragraphs[0].runs[0]
            value_run.font.color.rgb = self.colors['dark_gray']
    
    def _add_word_dashboard(self, doc, coverage_data):
        """Add visual dashboard to Word"""
        doc.add_heading('📊 Enterprise Coverage Dashboard', level=1)
        
        # Coverage Score
        score_para = doc.add_paragraph()
        score_para.add_run(f"Coverage Score: ").bold = True
        score_run = score_para.add_run(f"{coverage_data.get('overall_score', 0)}%")
        score_run.font.size = Pt(24)
        score_run.font.color.rgb = self._get_score_color_rgb(coverage_data.get('overall_score', 0))
        
        # Enterprise Readiness
        readiness_para = doc.add_paragraph()
        readiness_para.add_run(f"Enterprise Readiness: ").bold = True
        readiness_para.add_run(coverage_data.get('enterprise_readiness', 'Unknown'))
        
        doc.add_paragraph()
        
        # Coverage Summary Table
        summary_table = doc.add_table(rows=3, cols=2)
        summary_table.style = 'Light Grid Accent 1'
        
        present_count = len(coverage_data.get('coverage_analysis', {}).get('present_elements', []))
        missing_count = len(coverage_data.get('coverage_analysis', {}).get('missing_elements', []))
        
        summary_table.rows[0].cells[0].text = '✅ Covered Elements'
        summary_table.rows[0].cells[1].text = str(present_count)
        
        summary_table.rows[1].cells[0].text = '❌ Missing Elements'
        summary_table.rows[1].cells[1].text = str(missing_count)
        
        summary_table.rows[2].cells[0].text = '📈 Total Elements'
        summary_table.rows[2].cells[1].text = '10'
    
    def _add_word_executive_summary(self, doc, story_data, coverage_data):
        """Add executive summary section"""
        doc.add_heading('📋 Executive Summary', level=1)
        
        # Business Goal
        doc.add_heading('Business Objective', level=2)
        doc.add_paragraph(story_data.get('business_goal', 'Not specified'))
        
        # Key Stakeholder
        doc.add_heading('Key Stakeholder', level=2)
        doc.add_paragraph(story_data.get('actor', 'Not specified'))
        
        # Timeline Estimate
        doc.add_heading('Estimated Timeline', level=2)
        timeline = self._estimate_timeline(story_data)
        doc.add_paragraph(timeline)
        
        # Risk Summary
        if story_data.get('risks'):
            doc.add_heading('Key Risks', level=2)
            for risk in story_data['risks'][:3]:  # Top 3 risks
                doc.add_paragraph(f"• {risk}", style='List Bullet')
    
    def _add_word_content(self, doc, story_data, template):
        """Add main content based on template"""
        doc.add_heading('📝 Detailed Requirements', level=1)
        
        # All 10 enterprise elements
        sections = [
            ('🎯 Business Goal', 'business_goal', False),
            ('👤 Actor', 'actor', False),
            ('⚡ Trigger', 'trigger', False),
            ('📋 Preconditions', 'preconditions', True),
            ('🔄 Functional Flow', 'functional_flow', True),
            ('✅ Validations', 'validations', True),
            ('🎯 Acceptance Criteria', 'acceptance_criteria', True),
            ('🔒 Security', 'security', True),
            ('🔗 Dependencies', 'dependencies', True),
            ('⚠️ Risks', 'risks', True)
        ]
        
        for title, key, is_list in sections:
            if story_data.get(key):
                doc.add_heading(title, level=2)
                
                if is_list:
                    for item in story_data[key]:
                        doc.add_paragraph(f"• {item}", style='List Bullet')
                else:
                    doc.add_paragraph(story_data[key])
                
                doc.add_paragraph()  # Spacing
    
    def _export_pdf(self, story_data, template, coverage_data):
        """Export story to PDF with visual dashboard"""
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"enterprise_user_story_{timestamp}.pdf"
        filepath = os.path.join(self.temp_dir, filename)
        
        doc = SimpleDocTemplate(filepath, pagesize=letter)
        styles = getSampleStyleSheet()
        story = []
        
        # Custom styles
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=24,
            textColor=colors.HexColor('#667eea'),
            spaceAfter=30,
            alignment=1  # Center
        )
        
        # Cover Page
        self._add_pdf_cover_page(story, story_data, styles, title_style)
        story.append(PageBreak())
        
        # Dashboard
        if coverage_data:
            self._add_pdf_dashboard(story, coverage_data, styles)
            story.append(PageBreak())
        
        # Executive Summary
        self._add_pdf_executive_summary(story, story_data, coverage_data, styles)
        story.append(PageBreak())
        
        # Main Content
        self._add_pdf_content(story, story_data, styles)
        
        doc.build(story)
        return filepath
    
    def _add_pdf_cover_page(self, story, story_data, styles, title_style):
        """Add PDF cover page"""
        # Logo
        if os.path.exists(self.logo_path):
            logo = RLImage(self.logo_path, width=3*inch, height=1.2*inch)
            story.append(logo)
        
        story.append(Spacer(1, 50))
        
        # Title
        story.append(Paragraph("Enterprise User Story", title_style))
        story.append(Spacer(1, 20))
        
        # Business Goal
        if story_data.get('business_goal'):
            goal_text = story_data['business_goal'][:150]
            story.append(Paragraph(goal_text, styles['Normal']))
        
        story.append(Spacer(1, 50))
        
        # Metadata table
        meta_data = [
            ['Created Date:', datetime.now().strftime("%B %d, %Y")],
            ['Document Type:', 'Enterprise User Story'],
            ['Version:', '1.0']
        ]
        
        meta_table = Table(meta_data, colWidths=[2*inch, 3*inch])
        meta_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (0, -1), colors.HexColor('#f3f4f6')),
            ('TEXTCOLOR', (0, 0), (-1, -1), colors.black),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, -1), 12),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 12),
            ('GRID', (0, 0), (-1, -1), 1, colors.grey)
        ]))
        
        story.append(meta_table)
    
    def _add_pdf_dashboard(self, story, coverage_data, styles):
        """Add visual dashboard to PDF"""
        story.append(Paragraph("📊 Enterprise Coverage Dashboard", styles['Heading1']))
        story.append(Spacer(1, 20))
        
        # Coverage Score Circle
        score = coverage_data.get('overall_score', 0)
        score_drawing = self._create_score_circle(score)
        story.append(score_drawing)
        story.append(Spacer(1, 20))
        
        # Coverage Summary Table
        present_count = len(coverage_data.get('coverage_analysis', {}).get('present_elements', []))
        missing_count = len(coverage_data.get('coverage_analysis', {}).get('missing_elements', []))
        
        summary_data = [
            ['Metric', 'Value', 'Status'],
            ['Coverage Score', f"{score}%", self._get_status_emoji(score)],
            ['Covered Elements', str(present_count), '✅'],
            ['Missing Elements', str(missing_count), '❌'],
            ['Enterprise Readiness', coverage_data.get('enterprise_readiness', 'Unknown'), '📊']
        ]
        
        summary_table = Table(summary_data, colWidths=[2.5*inch, 1.5*inch, 1*inch])
        summary_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#667eea')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 14),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
            ('GRID', (0, 0), (-1, -1), 1, colors.black)
        ]))
        
        story.append(summary_table)
    
    def _create_score_circle(self, score):
        """Create visual score circle"""
        drawing = Drawing(200, 200)
        
        # Background circle
        drawing.add(Circle(100, 100, 80, fillColor=colors.HexColor('#f3f4f6'), strokeColor=colors.grey))
        
        # Progress circle
        color = self._get_score_color(score)
        pie = Pie()
        pie.x = 20
        pie.y = 20
        pie.width = 160
        pie.height = 160
        pie.data = [score, 100-score]
        pie.slices[0].fillColor = color
        pie.slices[1].fillColor = colors.HexColor('#e5e7eb')
        pie.slices[0].strokeColor = None
        pie.slices[1].strokeColor = None
        
        drawing.add(pie)
        
        # Score text
        score_text = String(100, 95, f"{score}%", textAnchor='middle', fontSize=24, fillColor=colors.black)
        drawing.add(score_text)
        
        return drawing
    
    def _add_pdf_executive_summary(self, story, story_data, coverage_data, styles):
        """Add executive summary to PDF"""
        story.append(Paragraph("📋 Executive Summary", styles['Heading1']))
        story.append(Spacer(1, 12))
        
        # Business Objective
        story.append(Paragraph("<b>Business Objective:</b>", styles['Heading2']))
        story.append(Paragraph(story_data.get('business_goal', 'Not specified'), styles['Normal']))
        story.append(Spacer(1, 12))
        
        # Key Stakeholder
        story.append(Paragraph("<b>Key Stakeholder:</b>", styles['Heading2']))
        story.append(Paragraph(story_data.get('actor', 'Not specified'), styles['Normal']))
        story.append(Spacer(1, 12))
        
        # Timeline
        story.append(Paragraph("<b>Estimated Timeline:</b>", styles['Heading2']))
        timeline = self._estimate_timeline(story_data)
        story.append(Paragraph(timeline, styles['Normal']))
        story.append(Spacer(1, 12))
        
        # Key Risks
        if story_data.get('risks'):
            story.append(Paragraph("<b>Key Risks:</b>", styles['Heading2']))
            for risk in story_data['risks'][:3]:
                story.append(Paragraph(f"• {risk}", styles['Normal']))
    
    def _add_pdf_content(self, story, story_data, styles):
        """Add main content to PDF"""
        story.append(Paragraph("📝 Detailed Requirements", styles['Heading1']))
        story.append(Spacer(1, 20))
        
        sections = [
            ('🎯 Business Goal', 'business_goal', False),
            ('👤 Actor', 'actor', False),
            ('⚡ Trigger', 'trigger', False),
            ('📋 Preconditions', 'preconditions', True),
            ('🔄 Functional Flow', 'functional_flow', True),
            ('✅ Validations', 'validations', True),
            ('🎯 Acceptance Criteria', 'acceptance_criteria', True),
            ('🔒 Security', 'security', True),
            ('🔗 Dependencies', 'dependencies', True),
            ('⚠️ Risks', 'risks', True)
        ]
        
        for title, key, is_list in sections:
            if story_data.get(key):
                story.append(Paragraph(title, styles['Heading2']))
                story.append(Spacer(1, 6))
                
                if is_list:
                    for item in story_data[key]:
                        story.append(Paragraph(f"• {item}", styles['Normal']))
                else:
                    story.append(Paragraph(story_data[key], styles['Normal']))
                
                story.append(Spacer(1, 12))
    
    def _export_png(self, story_data):
        """Export story to PNG image"""
        if not Image:
            raise ValueError("PIL/Pillow not properly installed. PNG export unavailable.")
        
        width, height = 800, 1400
        img = Image.new('RGB', (width, height), color='white')
        draw = ImageDraw.Draw(img)
        
        try:
            title_font = ImageFont.truetype("arial.ttf", 24)
            header_font = ImageFont.truetype("arial.ttf", 16)
            text_font = ImageFont.truetype("arial.ttf", 12)
        except:
            title_font = ImageFont.load_default()
            header_font = ImageFont.load_default()
            text_font = ImageFont.load_default()
        
        y_position = 20
        margin = 20
        
        if os.path.exists(self.logo_path):
            try:
                logo_img = Image.open(self.logo_path)
                logo_img = logo_img.resize((150, 45))
                img.paste(logo_img, (margin, y_position))
            except:
                pass
        
        draw.text((width - 300, y_position + 10), "Enterprise User Story", fill='black', font=title_font)
        y_position += 70
        
        if story_data.get('business_goal'):
            draw.text((margin, y_position), '🎯 Business Goal:', fill='black', font=header_font)
            y_position += 25
            goal_text = story_data['business_goal'][:80] + "..." if len(story_data['business_goal']) > 80 else story_data['business_goal']
            draw.text((margin, y_position), goal_text, fill='black', font=text_font)
            y_position += 35
        
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"enterprise_user_story_{timestamp}.png"
        filepath = os.path.join(self.temp_dir, filename)
        img.save(filepath)
        
        return filepath
    
    # Helper methods
    def _estimate_timeline(self, story_data):
        """Estimate development timeline"""
        base_weeks = 2
        
        if story_data.get('security') and len(story_data['security']) > 3:
            base_weeks += 1
        if story_data.get('dependencies') and len(story_data['dependencies']) > 3:
            base_weeks += 1
        if story_data.get('functional_flow') and len(story_data['functional_flow']) > 10:
            base_weeks += 1
        
        return f"{base_weeks}-{base_weeks + 1} weeks"
    
    def _get_score_color(self, score):
        """Get color based on score"""
        if score >= 90:
            return colors.HexColor('#10b981')
        elif score >= 70:
            return colors.HexColor('#f59e0b')
        elif score >= 50:
            return colors.HexColor('#f97316')
        else:
            return colors.HexColor('#ef4444')
    
    def _get_score_color_rgb(self, score):
        """Get RGB color for Word"""
        if score >= 90:
            return RGBColor(16, 185, 129)
        elif score >= 70:
            return RGBColor(245, 158, 11)
        elif score >= 50:
            return RGBColor(249, 115, 22)
        else:
            return RGBColor(239, 68, 68)
    
    def _get_status_emoji(self, score):
        """Get status emoji based on score"""
        if score >= 90:
            return '🟢'
        elif score >= 70:
            return '🟡'
        elif score >= 50:
            return '🟠'
        else:
            return '🔴'
