import os
import tempfile
import base64
import io
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.shared import OxmlElement, qn
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib import colors
from reportlab.platypus import Image as RLImage

class EnhancedStoryExporter:
    def __init__(self):
        self.temp_dir = tempfile.gettempdir()
        self.logo_path = os.path.join('static', 'images', 'anand_rathi_logo.png')
        
        # Corporate Color Scheme
        self.colors = {
            'primary': (30, 64, 175),
            'success_green': (5, 150, 105),
            'warning_orange': (217, 119, 6),
            'error_red': (220, 38, 38),
            'light_gray': (248, 250, 252),
            'secondary': (100, 116, 139),
            'text': (30, 41, 59),
            'accent_purple': (147, 51, 234)
        }
    
    def export_story(self, story_data, format_type, coverage_data=None, section_images=None):
        """Export story with enhanced styling"""
        if format_type == 'word':
            return self._export_word_corporate(story_data, coverage_data, section_images)
        elif format_type == 'pdf':
            return self._export_pdf_enhanced(story_data, coverage_data, section_images)
        else:
            raise ValueError(f"Unsupported format: {format_type}")
    
    def export_brd(self, brd_data, format_type, coverage_data=None, section_images=None):
        """Export BRD with enhanced styling"""
        if format_type == 'word':
            return self._export_brd_word_corporate(brd_data, coverage_data, section_images)
        elif format_type == 'pdf':
            return self._export_brd_pdf_enhanced(brd_data, coverage_data, section_images)
        else:
            raise ValueError(f"Unsupported format: {format_type}")
    
    def export_frd(self, frd_data, format_type, coverage_data=None, section_images=None):
        """Export FRD with enhanced styling"""
        if format_type == 'word':
            return self._export_frd_word_corporate(frd_data, coverage_data, section_images)
        elif format_type == 'pdf':
            return self._export_frd_pdf_enhanced(frd_data, coverage_data, section_images)
        else:
            raise ValueError(f"Unsupported format: {format_type}")
    
    def export_srd(self, srd_data, format_type, coverage_data=None, section_images=None):
        """Export SRD with enhanced styling"""
        if format_type == 'word':
            return self._export_srd_word_corporate(srd_data, coverage_data, section_images)
        elif format_type == 'pdf':
            return self._export_srd_pdf_enhanced(srd_data, coverage_data, section_images)
        else:
            raise ValueError(f"Unsupported format: {format_type}")
    
    def _export_word_corporate(self, story_data, coverage_data, section_images=None):
        """Export Word with corporate formatting"""
        doc = Document()
        
        # Set document margins and page setup
        self._setup_document_formatting(doc)
        
        # Add page borders
        self._add_page_borders(doc)
        
        # Add header and footer
        self._add_header_footer(doc)
        
        # Cover Page
        self._add_corporate_cover_page(doc, story_data)
        doc.add_page_break()
        
        # Dashboard
        if coverage_data:
            self._add_corporate_dashboard(doc, coverage_data)
            doc.add_page_break()
        
        # Content with corporate styling
        self._add_corporate_content(doc, story_data, section_images)
        
        # Save
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"Agile_Story_Document_{timestamp}.docx"
        filepath = os.path.join(self.temp_dir, filename)
        doc.save(filepath)
        
        return filepath
    
    def _setup_document_formatting(self, doc):
        """Setup document margins and formatting"""
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1.2)
            section.bottom_margin = Inches(1.0)
            section.left_margin = Inches(1.0)
            section.right_margin = Inches(1.0)
            section.page_width = Inches(8.5)
            section.page_height = Inches(11)
    
    def _add_page_borders(self, doc):
        """Add corporate page borders"""
        try:
            sections = doc.sections
            for section in sections:
                sectPr = section._sectPr
                pgBorders = OxmlElement('w:pgBorders')
                pgBorders.set(qn('w:offsetFrom'), 'page')
                
                # Add borders (top, bottom, left, right)
                for border_name in ['top', 'bottom', 'left', 'right']:
                    border = OxmlElement(f'w:{border_name}')
                    border.set(qn('w:val'), 'single')
                    border.set(qn('w:sz'), '8')  # 1pt
                    border.set(qn('w:space'), '10')  # 10pt margin
                    border.set(qn('w:color'), '404040')  # Dark gray
                    pgBorders.append(border)
                
                sectPr.append(pgBorders)
        except Exception:
            pass  # Skip if border fails
    
    def _add_header_footer(self, doc):
        """Add corporate header and footer"""
        try:
            section = doc.sections[0]
            
            # Header
            header = section.header
            header_para = header.paragraphs[0]
            header_para.clear()
            
            # Left: Company name
            left_run = header_para.add_run('ANAND RATHI IT PVT. LTD.')
            left_run.font.name = 'Calibri'
            left_run.font.size = Pt(10)
            left_run.font.bold = True
            left_run.font.color.rgb = RGBColor(*self.colors['primary'])
            
            # Center: Document title
            header_para.add_run('\t\tAgile Story Document')
            center_run = header_para.runs[-1]
            center_run.font.name = 'Calibri'
            center_run.font.size = Pt(10)
            center_run.font.bold = True
            
            # Right: Version
            header_para.add_run('\t\tVersion: 1.0')
            right_run = header_para.runs[-1]
            right_run.font.name = 'Calibri'
            right_run.font.size = Pt(10)
            
            # Add bottom border to header
            self._add_header_footer_border(header_para, 'bottom')
            
            # Footer
            footer = section.footer
            footer_para = footer.paragraphs[0]
            footer_para.clear()
            footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            footer_run = footer_para.add_run('Confidential – For Internal Use Only | Page ')
            footer_run.font.name = 'Calibri'
            footer_run.font.size = Pt(9)
            footer_run.font.color.rgb = RGBColor(*self.colors['secondary'])
            
            # Add top border to footer
            self._add_header_footer_border(footer_para, 'top')
            
        except Exception:
            pass  # Skip if header/footer fails
    
    def _add_header_footer_border(self, paragraph, position):
        """Add border to header/footer"""
        try:
            pPr = paragraph._element.get_or_add_pPr()
            pBdr = OxmlElement('w:pBdr')
            
            border = OxmlElement(f'w:{position}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '6')  # 0.75pt
            border.set(qn('w:space'), '1')
            border.set(qn('w:color'), '404040')
            
            pBdr.append(border)
            pPr.append(pBdr)
        except Exception:
            pass
    
    def _add_corporate_cover_page(self, doc, story_data):
        """Add corporate cover page"""
        # Logo
        if os.path.exists(self.logo_path):
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para.add_run()
            run.add_picture(self.logo_path, width=Inches(3.0))
        
        # Spacing
        for _ in range(3):
            doc.add_paragraph()
        
        # Title
        title = doc.add_heading('AGILE USER STORY DOCUMENT', level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title.runs[0]
        title_run.font.name = 'Calibri'
        title_run.font.size = Pt(20)
        title_run.font.color.rgb = RGBColor(*self.colors['primary'])
        title_run.bold = True
        
        # Subtitle
        if story_data.get('business_goal'):
            subtitle = doc.add_paragraph(story_data['business_goal'][:80])
            subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
            subtitle_run = subtitle.runs[0]
            subtitle_run.font.name = 'Calibri'
            subtitle_run.font.size = Pt(14)
            subtitle_run.font.color.rgb = RGBColor(*self.colors['secondary'])
            subtitle_run.italic = True
        
        # Spacing
        for _ in range(4):
            doc.add_paragraph()
        
        # Version History Table
        self._add_version_history_table(doc)
    
    def _add_version_history_table(self, doc):
        """Add version history table"""
        version_heading = doc.add_heading('Version History', level=2)
        self._apply_heading_style(version_heading, 2)
        
        table = doc.add_table(rows=2, cols=4)
        table.style = 'Table Grid'
        
        # Header row
        headers = ['Version', 'Date', 'Author', 'Changes']
        for i, header in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = header
            # Style header
            cell.paragraphs[0].runs[0].font.bold = True
            cell.paragraphs[0].runs[0].font.name = 'Calibri'
            cell.paragraphs[0].runs[0].font.size = Pt(11)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            # Add light gray background
            self._add_cell_shading(cell, 'E8E8E8')
        
        # Data row
        data = ['1.0', datetime.now().strftime('%Y-%m-%d'), 'Business Analyst', 'Initial version']
        for i, value in enumerate(data):
            cell = table.rows[1].cells[i]
            cell.text = value
            cell.paragraphs[0].runs[0].font.name = 'Calibri'
            cell.paragraphs[0].runs[0].font.size = Pt(10)
    
    def _add_cell_shading(self, cell, color):
        """Add background shading to table cell"""
        try:
            tcPr = cell._tc.get_or_add_tcPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'), 'clear')
            shd.set(qn('w:color'), 'auto')
            shd.set(qn('w:fill'), color)
            tcPr.append(shd)
        except Exception:
            pass
    
    def _add_corporate_dashboard(self, doc, coverage_data):
        """Add corporate dashboard"""
        # Objective heading
        obj_heading = doc.add_heading('OBJECTIVE', level=1)
        self._apply_heading_style(obj_heading, 1)
        
        # Business goal
        goal_para = doc.add_paragraph()
        goal_para.add_run('Business Goal: ').bold = True
        goal_para.add_run(coverage_data.get('business_goal', 'Enhance system functionality'))
        self._apply_body_style(goal_para)
        
        # Coverage metrics table
        metrics_heading = doc.add_heading('Coverage Metrics', level=2)
        self._apply_heading_style(metrics_heading, 2)
        
        metrics_table = doc.add_table(rows=4, cols=2)
        metrics_table.style = 'Table Grid'
        
        present = len(coverage_data.get('coverage_analysis', {}).get('present_elements', []))
        missing = len(coverage_data.get('coverage_analysis', {}).get('missing_elements', []))
        
        metrics_data = [
            ('Coverage Score', f"{coverage_data.get('overall_score', 0)}%"),
            ('Covered Elements', str(present)),
            ('Missing Elements', str(missing)),
            ('Enterprise Readiness', coverage_data.get('enterprise_readiness', 'Unknown'))
        ]
        
        for i, (label, value) in enumerate(metrics_data):
            row = metrics_table.rows[i]
            row.cells[0].text = label
            row.cells[1].text = value
            
            # Style cells
            for cell in row.cells:
                cell.paragraphs[0].runs[0].font.name = 'Calibri'
                cell.paragraphs[0].runs[0].font.size = Pt(11)
            
            row.cells[0].paragraphs[0].runs[0].font.bold = True
    
    def _add_corporate_content(self, doc, story_data, section_images=None):
        """Add corporate content sections"""
        # User Story Block
        story_heading = doc.add_heading('USER STORY', level=1)
        self._apply_heading_style(story_heading, 1)
        
        # Story details table
        story_table = doc.add_table(rows=4, cols=2)
        story_table.style = 'Table Grid'
        
        story_details = [
            ('Story ID:', 'US-001'),
            ('Story Name:', story_data.get('business_goal', 'Primary Feature')[:50]),
            ('Module:', 'Core Application'),
            ('Writer:', 'Business Analyst')
        ]
        
        for i, (label, value) in enumerate(story_details):
            row = story_table.rows[i]
            row.cells[0].text = label
            row.cells[1].text = value
            
            row.cells[0].paragraphs[0].runs[0].font.bold = True
            row.cells[0].paragraphs[0].runs[0].font.name = 'Calibri'
            row.cells[0].paragraphs[0].runs[0].font.size = Pt(11)
            row.cells[1].paragraphs[0].runs[0].font.name = 'Calibri'
            row.cells[1].paragraphs[0].runs[0].font.size = Pt(11)
        
        # Add horizontal line
        self._add_section_separator(doc)
        
        # Main sections with image support
        sections = [
            ('FUNCTIONAL FLOW', 'functional_flow', True, 'functional-flow'),
            ('VALIDATIONS', 'validations', True, 'validations'),
            ('ACCEPTANCE CRITERIA', 'acceptance_criteria', True, 'acceptance-criteria'),
            ('SECURITY REQUIREMENTS', 'security', True, 'security'),
            ('DEPENDENCIES', 'dependencies', True, 'dependencies'),
            ('RISKS & MITIGATION', 'risks', True, 'risks')
        ]
        
        for title, key, is_list, section_id in sections:
            if story_data.get(key):
                section_heading = doc.add_heading(title, level=1)
                self._apply_heading_style(section_heading, 1)
                
                if is_list:
                    for item in story_data[key]:
                        bullet_para = doc.add_paragraph(f'• {item}')
                        self._apply_body_style(bullet_para)
                        bullet_para.paragraph_format.left_indent = Inches(0.5)
                else:
                    content_para = doc.add_paragraph(story_data[key])
                    self._apply_body_style(content_para)
                
                # Add section images if available
                self._add_section_images(doc, section_id, section_images)
                
                self._add_section_separator(doc)
        
        # Non-Functional Requirements
        nfr_heading = doc.add_heading('NON-FUNCTIONAL REQUIREMENTS', level=1)
        self._apply_heading_style(nfr_heading, 1)
        
        nfr_items = [
            'Performance: Response time < 2 seconds',
            'Security: Authentication and authorization required',
            'Audit: All actions must be logged',
            'Logging: Comprehensive error and activity logging'
        ]
        
        for item in nfr_items:
            nfr_para = doc.add_paragraph(f'• {item}')
            self._apply_body_style(nfr_para)
            nfr_para.paragraph_format.left_indent = Inches(0.5)
    
    def _add_section_separator(self, doc):
        """Add horizontal line separator"""
        separator_para = doc.add_paragraph()
        separator_para.paragraph_format.space_before = Pt(14)
        separator_para.paragraph_format.space_after = Pt(10)
        separator_run = separator_para.add_run('_' * 80)
        separator_run.font.color.rgb = RGBColor(*self.colors['secondary'])
        separator_run.font.size = Pt(8)
        separator_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    def _apply_heading_style(self, heading, level):
        """Apply corporate heading styles"""
        heading_run = heading.runs[0]
        heading_run.font.name = 'Calibri'
        
        if level == 1:
            heading_run.font.size = Pt(16)
            heading_run.font.bold = True
            heading.paragraph_format.space_before = Pt(18)
            heading.paragraph_format.space_after = Pt(6)
            heading.paragraph_format.keep_with_next = True
        elif level == 2:
            heading_run.font.size = Pt(14)
            heading_run.font.bold = True
            heading.paragraph_format.space_before = Pt(12)
            heading.paragraph_format.space_after = Pt(4)
        elif level == 3:
            heading_run.font.size = Pt(12)
            heading_run.font.bold = True
            heading.paragraph_format.space_before = Pt(8)
            heading.paragraph_format.space_after = Pt(4)
        
        heading_run.font.color.rgb = RGBColor(*self.colors['primary'])
    
    def _apply_body_style(self, paragraph):
        """Apply corporate body text styles"""
        for run in paragraph.runs:
            run.font.name = 'Calibri'
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(*self.colors['text'])
        
        paragraph.paragraph_format.line_spacing = 1.15
        paragraph.paragraph_format.space_after = Pt(6)
    
    def _export_pdf_enhanced(self, story_data, coverage_data, section_images=None):
        """PDF with enhanced styling"""
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"Agile_Story_Document_{timestamp}.pdf"
        filepath = os.path.join(self.temp_dir, filename)
        
        doc = SimpleDocTemplate(filepath, pagesize=letter)
        styles = getSampleStyleSheet()
        story = []
        
        # Custom styles
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=20,
            textColor=colors.HexColor('#1e40af'),
            spaceAfter=30,
            alignment=1,
            fontName='Helvetica-Bold'
        )
        
        # Cover page
        if os.path.exists(self.logo_path):
            logo = RLImage(self.logo_path, width=3*inch, height=1.2*inch)
            story.append(logo)
        
        story.append(Spacer(1, 50))
        story.append(Paragraph("AGILE USER STORY DOCUMENT", title_style))
        story.append(Spacer(1, 20))
        
        if story_data.get('business_goal'):
            story.append(Paragraph(story_data['business_goal'][:150], styles['Normal']))
        
        story.append(PageBreak())
        
        # Dashboard
        if coverage_data:
            self._add_pdf_corporate_dashboard(story, coverage_data, styles)
            story.append(PageBreak())
        
        # Content
        self._add_pdf_corporate_content(story, story_data, styles)
        
        doc.build(story)
        return filepath
    
    def _add_pdf_corporate_dashboard(self, story, coverage_data, styles):
        """PDF corporate dashboard"""
        story.append(Paragraph("OBJECTIVE", styles['Heading1']))
        story.append(Spacer(1, 12))
        
        # Business goal
        story.append(Paragraph(f"<b>Business Goal:</b> {coverage_data.get('business_goal', 'Enhance system functionality')}", styles['Normal']))
        story.append(Spacer(1, 20))
        
        # Coverage metrics
        story.append(Paragraph("Coverage Metrics", styles['Heading2']))
        story.append(Spacer(1, 12))
        
        present = len(coverage_data.get('coverage_analysis', {}).get('present_elements', []))
        missing = len(coverage_data.get('coverage_analysis', {}).get('missing_elements', []))
        
        metrics_data = [
            ['Metric', 'Value'],
            ['Coverage Score', f"{coverage_data.get('overall_score', 0)}%"],
            ['Covered Elements', str(present)],
            ['Missing Elements', str(missing)],
            ['Enterprise Readiness', coverage_data.get('enterprise_readiness', 'Unknown')]
        ]
        
        metrics_table = Table(metrics_data, colWidths=[3*inch, 2*inch])
        metrics_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#E8E8E8')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.black),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, -1), 11),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 12),
            ('GRID', (0, 0), (-1, -1), 0.75, colors.black)
        ]))
        
        story.append(metrics_table)
        story.append(Spacer(1, 20))
    
    def _add_pdf_corporate_content(self, story, story_data, styles):
        """PDF corporate content"""
        # User Story section
        story.append(Paragraph("USER STORY", styles['Heading1']))
        story.append(Spacer(1, 12))
        
        # Story details
        story_details = [
            ['Field', 'Value'],
            ['Story ID', 'US-001'],
            ['Story Name', story_data.get('business_goal', 'Primary Feature')[:50]],
            ['Module', 'Core Application'],
            ['Writer', 'Business Analyst']
        ]
        
        story_table = Table(story_details, colWidths=[2*inch, 4*inch])
        story_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#E8E8E8')),
            ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, -1), 11),
            ('GRID', (0, 0), (-1, -1), 0.75, colors.black),
            ('LEFTPADDING', (0, 0), (-1, -1), 12),
            ('RIGHTPADDING', (0, 0), (-1, -1), 12),
            ('TOPPADDING', (0, 0), (-1, -1), 8),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 8)
        ]))
        
        story.append(story_table)
        story.append(Spacer(1, 20))
        
        # Main sections
        sections = [
            ('FUNCTIONAL FLOW', 'functional_flow', True),
            ('VALIDATIONS', 'validations', True),
            ('ACCEPTANCE CRITERIA', 'acceptance_criteria', True),
            ('SECURITY REQUIREMENTS', 'security', True),
            ('DEPENDENCIES', 'dependencies', True),
            ('RISKS & MITIGATION', 'risks', True)
        ]
        
        for title, key, is_list in sections:
            if story_data.get(key):
                story.append(Paragraph(title, styles['Heading1']))
                story.append(Spacer(1, 6))
                
                if is_list:
                    for item in story_data[key]:
                        story.append(Paragraph(f"• {item}", styles['Normal']))
                else:
                    story.append(Paragraph(story_data[key], styles['Normal']))
                
                story.append(Spacer(1, 16))
    
    def _export_brd_word_corporate(self, brd_data, coverage_data, section_images=None):
        """Export BRD Word with corporate formatting"""
        doc = Document()
        
        # Set document margins and page setup
        self._setup_document_formatting(doc)
        
        # Add page borders
        self._add_page_borders(doc)
        
        # Add header and footer for BRD
        self._add_brd_header_footer(doc)
        
        # Cover Page
        self._add_brd_corporate_cover_page(doc, brd_data)
        doc.add_page_break()
        
        # Document Control
        self._add_brd_document_control(doc)
        doc.add_page_break()
        
        # Dashboard
        if coverage_data:
            self._add_brd_corporate_dashboard(doc, coverage_data)
            doc.add_page_break()
        
        # BRD Content with corporate styling
        self._add_brd_corporate_content(doc, brd_data, section_images)
        
        # Save
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"Business_Requirements_Document_{timestamp}.docx"
        filepath = os.path.join(self.temp_dir, filename)
        doc.save(filepath)
        
        return filepath
    
    def _add_brd_header_footer(self, doc):
        """Add BRD corporate header and footer"""
        try:
            section = doc.sections[0]
            
            # Header
            header = section.header
            header_para = header.paragraphs[0]
            header_para.clear()
            
            # Left: Company name
            left_run = header_para.add_run('ANAND RATHI IT PVT. LTD.')
            left_run.font.name = 'Calibri'
            left_run.font.size = Pt(10)
            left_run.font.bold = True
            left_run.font.color.rgb = RGBColor(*self.colors['primary'])
            
            # Center: Document title
            header_para.add_run('\t\tBusiness Requirements Document')
            center_run = header_para.runs[-1]
            center_run.font.name = 'Calibri'
            center_run.font.size = Pt(10)
            center_run.font.bold = True
            
            # Right: Version
            header_para.add_run('\t\tVersion: 1.0')
            right_run = header_para.runs[-1]
            right_run.font.name = 'Calibri'
            right_run.font.size = Pt(10)
            
            # Add bottom border to header
            self._add_header_footer_border(header_para, 'bottom')
            
            # Footer
            footer = section.footer
            footer_para = footer.paragraphs[0]
            footer_para.clear()
            footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            footer_run = footer_para.add_run('Confidential – For Internal Use Only | Page ')
            footer_run.font.name = 'Calibri'
            footer_run.font.size = Pt(9)
            footer_run.font.color.rgb = RGBColor(*self.colors['secondary'])
            
            # Add top border to footer
            self._add_header_footer_border(footer_para, 'top')
            
        except Exception:
            pass  # Skip if header/footer fails
    
    def _add_brd_corporate_cover_page(self, doc, brd_data):
        """Add BRD corporate cover page"""
        # Logo
        if os.path.exists(self.logo_path):
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para.add_run()
            run.add_picture(self.logo_path, width=Inches(3.0))
        
        # Spacing
        for _ in range(3):
            doc.add_paragraph()
        
        # Title
        title = doc.add_heading('BUSINESS REQUIREMENTS DOCUMENT', level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title.runs[0]
        title_run.font.name = 'Calibri'
        title_run.font.size = Pt(20)
        title_run.font.color.rgb = RGBColor(*self.colors['primary'])
        title_run.bold = True
        
        # Project Name
        if brd_data.get('project_name'):
            subtitle = doc.add_paragraph(brd_data['project_name'])
            subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
            subtitle_run = subtitle.runs[0]
            subtitle_run.font.name = 'Calibri'
            subtitle_run.font.size = Pt(14)
            subtitle_run.font.color.rgb = RGBColor(*self.colors['secondary'])
            subtitle_run.italic = True
        
        # Spacing
        for _ in range(4):
            doc.add_paragraph()
        
        # Version History Table
        self._add_version_history_table(doc)
    
    def _add_brd_document_control(self, doc):
        """Add BRD document control section"""
        control_heading = doc.add_heading('DOCUMENT CONTROL', level=1)
        self._apply_heading_style(control_heading, 1)
        
        # Document Info Table
        info_table = doc.add_table(rows=6, cols=2)
        info_table.style = 'Table Grid'
        
        info_data = [
            ('Document Title', 'Business Requirements Document'),
            ('Document ID', 'BRD-001'),
            ('Version', '1.0'),
            ('Author', 'Business Analyst'),
            ('Department', 'Business Analysis'),
            ('Date', datetime.now().strftime('%Y-%m-%d'))
        ]
        
        for i, (label, value) in enumerate(info_data):
            row = info_table.rows[i]
            row.cells[0].text = label
            row.cells[1].text = value
            
            # Style cells
            row.cells[0].paragraphs[0].runs[0].font.bold = True
            row.cells[0].paragraphs[0].runs[0].font.name = 'Calibri'
            row.cells[0].paragraphs[0].runs[0].font.size = Pt(11)
            row.cells[1].paragraphs[0].runs[0].font.name = 'Calibri'
            row.cells[1].paragraphs[0].runs[0].font.size = Pt(11)
            
            # Add light gray background to header column
            self._add_cell_shading(row.cells[0], 'E8E8E8')
    
    def _add_brd_corporate_dashboard(self, doc, coverage_data):
        """Add BRD corporate dashboard"""
        # Coverage metrics heading
        metrics_heading = doc.add_heading('BRD COVERAGE ANALYSIS', level=1)
        self._apply_heading_style(metrics_heading, 1)
        
        # Project name
        if coverage_data.get('project_name'):
            project_para = doc.add_paragraph()
            project_para.add_run('Project: ').bold = True
            project_para.add_run(coverage_data.get('project_name', 'Business Requirements Project'))
            self._apply_body_style(project_para)
        
        # Coverage metrics table
        metrics_table = doc.add_table(rows=4, cols=2)
        metrics_table.style = 'Table Grid'
        
        present = len(coverage_data.get('coverage_analysis', {}).get('present_elements', []))
        missing = len(coverage_data.get('coverage_analysis', {}).get('missing_elements', []))
        
        metrics_data = [
            ('Coverage Score', f"{coverage_data.get('overall_score', 0)}%"),
            ('Covered Elements', str(present)),
            ('Missing Elements', str(missing)),
            ('Enterprise Readiness', coverage_data.get('enterprise_readiness', 'Unknown'))
        ]
        
        for i, (label, value) in enumerate(metrics_data):
            row = metrics_table.rows[i]
            row.cells[0].text = label
            row.cells[1].text = value
            
            # Style cells
            for cell in row.cells:
                cell.paragraphs[0].runs[0].font.name = 'Calibri'
                cell.paragraphs[0].runs[0].font.size = Pt(11)
            
            row.cells[0].paragraphs[0].runs[0].font.bold = True
            self._add_cell_shading(row.cells[0], 'E8E8E8')
    
    def _add_brd_corporate_content(self, doc, brd_data, section_images=None):
        """Add BRD corporate content sections"""
        # Executive Summary
        if brd_data.get('executive_summary'):
            exec_heading = doc.add_heading('EXECUTIVE SUMMARY', level=1)
            self._apply_heading_style(exec_heading, 1)
            
            exec_summary = brd_data['executive_summary']
            if exec_summary.get('background'):
                bg_para = doc.add_paragraph()
                bg_para.add_run('Background: ').bold = True
                bg_para.add_run(exec_summary['background'])
                self._apply_body_style(bg_para)
            
            if exec_summary.get('problem_statement'):
                prob_para = doc.add_paragraph()
                prob_para.add_run('Problem Statement: ').bold = True
                prob_para.add_run(exec_summary['problem_statement'])
                self._apply_body_style(prob_para)
            
            if exec_summary.get('business_need'):
                need_para = doc.add_paragraph()
                need_para.add_run('Business Need: ').bold = True
                need_para.add_run(exec_summary['business_need'])
                self._apply_body_style(need_para)
            
            # Add section images
            self._add_section_images(doc, 'executive-summary', section_images)
            self._add_section_separator(doc)
        
        # Business Objectives
        if brd_data.get('business_objectives'):
            obj_heading = doc.add_heading('BUSINESS OBJECTIVES', level=1)
            self._apply_heading_style(obj_heading, 1)
            
            for i, obj in enumerate(brd_data['business_objectives'], 1):
                obj_para = doc.add_paragraph(f"{i}. {obj.get('objective', 'Objective')}")
                self._apply_body_style(obj_para)
                if obj.get('kpi'):
                    kpi_para = doc.add_paragraph(f"   KPI: {obj['kpi']}")
                    self._apply_body_style(kpi_para)
                    kpi_para.paragraph_format.left_indent = Inches(0.5)
            
            # Add section images
            self._add_section_images(doc, 'business-objectives', section_images)
            self._add_section_separator(doc)
        
        # Scope
        if brd_data.get('scope'):
            scope_heading = doc.add_heading('PROJECT SCOPE', level=1)
            self._apply_heading_style(scope_heading, 1)
            
            scope = brd_data['scope']
            if scope.get('in_scope'):
                in_scope_heading = doc.add_heading('In-Scope', level=2)
                self._apply_heading_style(in_scope_heading, 2)
                for item in scope['in_scope']:
                    bullet_para = doc.add_paragraph(f'• {item}')
                    self._apply_body_style(bullet_para)
                    bullet_para.paragraph_format.left_indent = Inches(0.5)
            
            if scope.get('out_of_scope'):
                out_scope_heading = doc.add_heading('Out-of-Scope', level=2)
                self._apply_heading_style(out_scope_heading, 2)
                for item in scope['out_of_scope']:
                    bullet_para = doc.add_paragraph(f'• {item}')
                    self._apply_body_style(bullet_para)
                    bullet_para.paragraph_format.left_indent = Inches(0.5)
            
            # Add section images
            self._add_section_images(doc, 'scope', section_images)
            self._add_section_separator(doc)
        
        # Stakeholders
        if brd_data.get('stakeholders'):
            stakeholder_heading = doc.add_heading('STAKEHOLDER LIST', level=1)
            self._apply_heading_style(stakeholder_heading, 1)
            
            # Stakeholder table
            stakeholder_table = doc.add_table(rows=len(brd_data['stakeholders']) + 1, cols=4)
            stakeholder_table.style = 'Table Grid'
            
            # Header row
            headers = ['Name', 'Role', 'Department', 'Responsibilities']
            for i, header in enumerate(headers):
                cell = stakeholder_table.rows[0].cells[i]
                cell.text = header
                cell.paragraphs[0].runs[0].font.bold = True
                cell.paragraphs[0].runs[0].font.name = 'Calibri'
                cell.paragraphs[0].runs[0].font.size = Pt(11)
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                self._add_cell_shading(cell, 'E8E8E8')
            
            # Data rows
            for i, stakeholder in enumerate(brd_data['stakeholders'], 1):
                row = stakeholder_table.rows[i]
                row.cells[0].text = stakeholder.get('name', '')
                row.cells[1].text = stakeholder.get('role', '')
                row.cells[2].text = stakeholder.get('department', '')
                row.cells[3].text = stakeholder.get('responsibilities', '')
                
                for cell in row.cells:
                    cell.paragraphs[0].runs[0].font.name = 'Calibri'
                    cell.paragraphs[0].runs[0].font.size = Pt(10)
            
            # Add section images
            self._add_section_images(doc, 'stakeholders', section_images)
            self._add_section_separator(doc)
        
        # Business Requirements
        if brd_data.get('business_requirements'):
            req_heading = doc.add_heading('BUSINESS REQUIREMENTS', level=1)
            self._apply_heading_style(req_heading, 1)
            
            # Requirements table
            req_table = doc.add_table(rows=len(brd_data['business_requirements']) + 1, cols=6)
            req_table.style = 'Table Grid'
            
            # Header row
            headers = ['BR ID', 'Title', 'Description', 'Priority', 'Source', 'Acceptance Criteria']
            for i, header in enumerate(headers):
                cell = req_table.rows[0].cells[i]
                cell.text = header
                cell.paragraphs[0].runs[0].font.bold = True
                cell.paragraphs[0].runs[0].font.name = 'Calibri'
                cell.paragraphs[0].runs[0].font.size = Pt(10)
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                self._add_cell_shading(cell, 'E8E8E8')
            
            # Data rows
            for i, req in enumerate(brd_data['business_requirements'], 1):
                row = req_table.rows[i]
                row.cells[0].text = req.get('br_id', '')
                row.cells[1].text = req.get('title', '')
                row.cells[2].text = req.get('description', '')
                row.cells[3].text = req.get('priority', '')
                row.cells[4].text = req.get('source', '')
                row.cells[5].text = req.get('acceptance_criteria', '')
                
                for cell in row.cells:
                    cell.paragraphs[0].runs[0].font.name = 'Calibri'
                    cell.paragraphs[0].runs[0].font.size = Pt(9)
            
            # Add section images
            self._add_section_images(doc, 'business-requirements', section_images)
            self._add_section_separator(doc)
        
        # Risks
        if brd_data.get('risks'):
            risk_heading = doc.add_heading('RISKS & MITIGATION', level=1)
            self._apply_heading_style(risk_heading, 1)
            
            # Risk table
            risk_table = doc.add_table(rows=len(brd_data['risks']) + 1, cols=5)
            risk_table.style = 'Table Grid'
            
            # Header row
            headers = ['Risk ID', 'Description', 'Impact', 'Likelihood', 'Mitigation Plan']
            for i, header in enumerate(headers):
                cell = risk_table.rows[0].cells[i]
                cell.text = header
                cell.paragraphs[0].runs[0].font.bold = True
                cell.paragraphs[0].runs[0].font.name = 'Calibri'
                cell.paragraphs[0].runs[0].font.size = Pt(11)
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                self._add_cell_shading(cell, 'E8E8E8')
            
            # Data rows
            for i, risk in enumerate(brd_data['risks'], 1):
                row = risk_table.rows[i]
                row.cells[0].text = risk.get('risk_id', '')
                row.cells[1].text = risk.get('description', '')
                row.cells[2].text = risk.get('impact', '')
                row.cells[3].text = risk.get('likelihood', '')
                row.cells[4].text = risk.get('mitigation', '')
                
                for cell in row.cells:
                    cell.paragraphs[0].runs[0].font.name = 'Calibri'
                    cell.paragraphs[0].runs[0].font.size = Pt(10)
            
            # Add section images
            self._add_section_images(doc, 'risks', section_images)
        
        # Approval Section
        approval_heading = doc.add_heading('APPROVAL SECTION', level=1)
        self._apply_heading_style(approval_heading, 1)
        
        approval_table = doc.add_table(rows=4, cols=3)
        approval_table.style = 'Table Grid'
        
        # Header row
        headers = ['Role', 'Name', 'Signature & Date']
        for i, header in enumerate(headers):
            cell = approval_table.rows[0].cells[i]
            cell.text = header
            cell.paragraphs[0].runs[0].font.bold = True
            cell.paragraphs[0].runs[0].font.name = 'Calibri'
            cell.paragraphs[0].runs[0].font.size = Pt(11)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            self._add_cell_shading(cell, 'E8E8E8')
        
        # Approval rows
        approval_roles = ['Business Owner', 'Project Manager', 'IT Manager']
        for i, role in enumerate(approval_roles, 1):
            row = approval_table.rows[i]
            row.cells[0].text = role
            row.cells[1].text = ''
            row.cells[2].text = ''
            
            for cell in row.cells:
                cell.paragraphs[0].runs[0].font.name = 'Calibri'
                cell.paragraphs[0].runs[0].font.size = Pt(11)
    
    def _export_brd_pdf_enhanced(self, brd_data, coverage_data, section_images=None):
        """Export BRD PDF with enhanced styling"""
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"Business_Requirements_Document_{timestamp}.pdf"
        filepath = os.path.join(self.temp_dir, filename)
        
        doc = SimpleDocTemplate(filepath, pagesize=letter)
        styles = getSampleStyleSheet()
        story = []
        
        # Custom styles
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=20,
            textColor=colors.HexColor('#1e40af'),
            spaceAfter=30,
            alignment=1,
            fontName='Helvetica-Bold'
        )
        
        # Cover page
        if os.path.exists(self.logo_path):
            logo = RLImage(self.logo_path, width=3*inch, height=1.2*inch)
            story.append(logo)
        
        story.append(Spacer(1, 50))
        story.append(Paragraph("BUSINESS REQUIREMENTS DOCUMENT", title_style))
        story.append(Spacer(1, 20))
        
        if brd_data.get('project_name'):
            story.append(Paragraph(brd_data['project_name'], styles['Normal']))
        
        story.append(PageBreak())
        
        # Dashboard
        if coverage_data:
            self._add_brd_pdf_corporate_dashboard(story, coverage_data, styles)
            story.append(PageBreak())
        
        # Content
        self._add_brd_pdf_corporate_content(story, brd_data, styles)
        
        doc.build(story)
        return filepath
    
    def _add_brd_pdf_corporate_dashboard(self, story, coverage_data, styles):
        """BRD PDF corporate dashboard"""
        story.append(Paragraph("BRD COVERAGE ANALYSIS", styles['Heading1']))
        story.append(Spacer(1, 12))
        
        # Project name
        if coverage_data.get('project_name'):
            story.append(Paragraph(f"<b>Project:</b> {coverage_data.get('project_name', 'Business Requirements Project')}", styles['Normal']))
            story.append(Spacer(1, 20))
        
        # Coverage metrics
        present = len(coverage_data.get('coverage_analysis', {}).get('present_elements', []))
        missing = len(coverage_data.get('coverage_analysis', {}).get('missing_elements', []))
        
        metrics_data = [
            ['Metric', 'Value'],
            ['Coverage Score', f"{coverage_data.get('overall_score', 0)}%"],
            ['Covered Elements', str(present)],
            ['Missing Elements', str(missing)],
            ['Enterprise Readiness', coverage_data.get('enterprise_readiness', 'Unknown')]
        ]
        
        metrics_table = Table(metrics_data, colWidths=[3*inch, 2*inch])
        metrics_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#E8E8E8')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.black),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, -1), 11),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 12),
            ('GRID', (0, 0), (-1, -1), 0.75, colors.black)
        ]))
        
        story.append(metrics_table)
        story.append(Spacer(1, 20))
    
    def _add_brd_pdf_corporate_content(self, story, brd_data, styles):
        """BRD PDF corporate content"""
        # Executive Summary
        if brd_data.get('executive_summary'):
            story.append(Paragraph("EXECUTIVE SUMMARY", styles['Heading1']))
            story.append(Spacer(1, 12))
            
            exec_summary = brd_data['executive_summary']
            if exec_summary.get('background'):
                story.append(Paragraph(f"<b>Background:</b> {exec_summary['background']}", styles['Normal']))
            if exec_summary.get('problem_statement'):
                story.append(Paragraph(f"<b>Problem Statement:</b> {exec_summary['problem_statement']}", styles['Normal']))
            if exec_summary.get('business_need'):
                story.append(Paragraph(f"<b>Business Need:</b> {exec_summary['business_need']}", styles['Normal']))
            
            story.append(Spacer(1, 20))
        
        # Business Objectives
        if brd_data.get('business_objectives'):
            story.append(Paragraph("BUSINESS OBJECTIVES", styles['Heading1']))
            story.append(Spacer(1, 6))
            
            for i, obj in enumerate(brd_data['business_objectives'], 1):
                story.append(Paragraph(f"{i}. {obj.get('objective', 'Objective')}", styles['Normal']))
                if obj.get('kpi'):
                    story.append(Paragraph(f"   KPI: {obj['kpi']}", styles['Normal']))
            
            story.append(Spacer(1, 16))
    
    def _export_frd_word_corporate(self, frd_data, coverage_data, section_images=None):
        """Export FRD Word with corporate formatting"""
        doc = Document()
        
        # Set document margins and page setup
        self._setup_document_formatting(doc)
        
        # Add page borders
        self._add_page_borders(doc)
        
        # Add header and footer for FRD
        self._add_frd_header_footer(doc)
        
        # Cover Page
        self._add_frd_corporate_cover_page(doc, frd_data)
        doc.add_page_break()
        
        # Document Control
        self._add_frd_document_control(doc)
        doc.add_page_break()
        
        # Dashboard
        if coverage_data:
            self._add_frd_corporate_dashboard(doc, coverage_data)
            doc.add_page_break()
        
        # FRD Content with corporate styling
        self._add_frd_corporate_content(doc, frd_data, section_images)
        
        # Save
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"Functional_Requirements_Document_{timestamp}.docx"
        filepath = os.path.join(self.temp_dir, filename)
        doc.save(filepath)
        
        return filepath
    
    def _add_frd_header_footer(self, doc):
        """Add FRD corporate header and footer"""
        try:
            section = doc.sections[0]
            
            # Header
            header = section.header
            header_para = header.paragraphs[0]
            header_para.clear()
            
            # Left: Company name
            left_run = header_para.add_run('ANAND RATHI IT PVT. LTD.')
            left_run.font.name = 'Calibri'
            left_run.font.size = Pt(10)
            left_run.font.bold = True
            left_run.font.color.rgb = RGBColor(*self.colors['primary'])
            
            # Center: Document title
            header_para.add_run('\t\tFunctional Requirements Document')
            center_run = header_para.runs[-1]
            center_run.font.name = 'Calibri'
            center_run.font.size = Pt(10)
            center_run.font.bold = True
            
            # Right: Version
            header_para.add_run('\t\tVersion: 1.0')
            right_run = header_para.runs[-1]
            right_run.font.name = 'Calibri'
            right_run.font.size = Pt(10)
            
            # Add bottom border to header
            self._add_header_footer_border(header_para, 'bottom')
            
            # Footer
            footer = section.footer
            footer_para = footer.paragraphs[0]
            footer_para.clear()
            footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            footer_run = footer_para.add_run('Confidential – For Internal Use Only | Page ')
            footer_run.font.name = 'Calibri'
            footer_run.font.size = Pt(9)
            footer_run.font.color.rgb = RGBColor(*self.colors['secondary'])
            
            # Add top border to footer
            self._add_header_footer_border(footer_para, 'top')
            
        except Exception:
            pass  # Skip if header/footer fails
    def _add_frd_header_footer(self, doc):
        """Add FRD corporate header and footer"""
        try:
            section = doc.sections[0]
            
            # Header
            header = section.header
            header_para = header.paragraphs[0]
            header_para.clear()
            
            # Left: Company name
            left_run = header_para.add_run('ANAND RATHI IT PVT. LTD.')
            left_run.font.name = 'Calibri'
            left_run.font.size = Pt(10)
            left_run.font.bold = True
            left_run.font.color.rgb = RGBColor(*self.colors['primary'])
            
            # Center: Document title
            header_para.add_run('\t\tFunctional Requirements Document')
            center_run = header_para.runs[-1]
            center_run.font.name = 'Calibri'
            center_run.font.size = Pt(10)
            center_run.font.bold = True
            
            # Right: Version
            header_para.add_run('\t\tVersion: 1.0')
            right_run = header_para.runs[-1]
            right_run.font.name = 'Calibri'
            right_run.font.size = Pt(10)
            
            # Add bottom border to header
            self._add_header_footer_border(header_para, 'bottom')
            
            # Footer
            footer = section.footer
            footer_para = footer.paragraphs[0]
            footer_para.clear()
            footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            footer_run = footer_para.add_run('Confidential – For Internal Use Only | Page ')
            footer_run.font.name = 'Calibri'
            footer_run.font.size = Pt(9)
            footer_run.font.color.rgb = RGBColor(*self.colors['secondary'])
            
            # Add top border to footer
            self._add_header_footer_border(footer_para, 'top')
            
        except Exception:
            pass  # Skip if FRD header/footer fails
    
    def _add_frd_corporate_cover_page(self, doc, frd_data):
        """Add FRD corporate cover page"""
        # Logo
        if os.path.exists(self.logo_path):
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para.add_run()
            run.add_picture(self.logo_path, width=Inches(3.0))
        
        # Spacing
        for _ in range(3):
            doc.add_paragraph()
        
        # Title
        title = doc.add_heading('FUNCTIONAL REQUIREMENTS DOCUMENT', level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title.runs[0]
        title_run.font.name = 'Calibri'
        title_run.font.size = Pt(20)
        title_run.font.color.rgb = RGBColor(*self.colors['primary'])
        title_run.bold = True
        
        # System Overview
        if frd_data.get('system_overview', {}).get('architecture'):
            subtitle = doc.add_paragraph(frd_data['system_overview']['architecture'][:80])
            subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
            subtitle_run = subtitle.runs[0]
            subtitle_run.font.name = 'Calibri'
            subtitle_run.font.size = Pt(14)
            subtitle_run.font.color.rgb = RGBColor(*self.colors['secondary'])
            subtitle_run.italic = True
        
        # Spacing
        for _ in range(4):
            doc.add_paragraph()
        
        # Version History Table
        self._add_version_history_table(doc)
    
    def _add_frd_document_control(self, doc):
        """Add FRD document control section"""
        control_heading = doc.add_heading('DOCUMENT CONTROL', level=1)
        self._apply_heading_style(control_heading, 1)
        
        # Document Info Table
        info_table = doc.add_table(rows=6, cols=2)
        info_table.style = 'Table Grid'
        
        info_data = [
            ('Document Title', 'Functional Requirements Document'),
            ('Document ID', 'FRD-001'),
            ('Version', '1.0'),
            ('Author', 'Technical Analyst'),
            ('Department', 'Technical Analysis'),
            ('Date', datetime.now().strftime('%Y-%m-%d'))
        ]
        
        for i, (label, value) in enumerate(info_data):
            row = info_table.rows[i]
            row.cells[0].text = label
            row.cells[1].text = value
            
            # Style cells
            row.cells[0].paragraphs[0].runs[0].font.bold = True
            row.cells[0].paragraphs[0].runs[0].font.name = 'Calibri'
            row.cells[0].paragraphs[0].runs[0].font.size = Pt(11)
            row.cells[1].paragraphs[0].runs[0].font.name = 'Calibri'
            row.cells[1].paragraphs[0].runs[0].font.size = Pt(11)
            
            # Add light gray background to header column
            self._add_cell_shading(row.cells[0], 'E8E8E8')
    
    def _add_frd_corporate_dashboard(self, doc, coverage_data):
        """Add FRD corporate dashboard"""
        # Coverage metrics heading
        metrics_heading = doc.add_heading('FRD COVERAGE ANALYSIS', level=1)
        self._apply_heading_style(metrics_heading, 1)
        
        # System complexity
        if coverage_data.get('system_complexity'):
            complexity_para = doc.add_paragraph()
            complexity_para.add_run('System Complexity: ').bold = True
            complexity_para.add_run(coverage_data.get('system_complexity', 'Medium'))
            self._apply_body_style(complexity_para)
        
        # Coverage metrics table
        metrics_table = doc.add_table(rows=4, cols=2)
        metrics_table.style = 'Table Grid'
        
        present = len(coverage_data.get('coverage_analysis', {}).get('present_elements', []))
        missing = len(coverage_data.get('coverage_analysis', {}).get('missing_elements', []))
        
        metrics_data = [
            ('Coverage Score', f"{coverage_data.get('overall_score', 0)}%"),
            ('Covered Elements', str(present)),
            ('Missing Elements', str(missing)),
            ('Enterprise Readiness', coverage_data.get('enterprise_readiness', 'Unknown'))
        ]
        
        for i, (label, value) in enumerate(metrics_data):
            row = metrics_table.rows[i]
            row.cells[0].text = label
            row.cells[1].text = value
            
            # Style cells
            for cell in row.cells:
                cell.paragraphs[0].runs[0].font.name = 'Calibri'
                cell.paragraphs[0].runs[0].font.size = Pt(11)
            
            row.cells[0].paragraphs[0].runs[0].font.bold = True
            self._add_cell_shading(row.cells[0], 'E8E8E8')
    
    def _add_frd_corporate_content(self, doc, frd_data, section_images=None):
        """Add FRD corporate content sections"""
        # System Overview
        if frd_data.get('system_overview'):
            sys_heading = doc.add_heading('SYSTEM OVERVIEW', level=1)
            self._apply_heading_style(sys_heading, 1)
            
            sys_overview = frd_data['system_overview']
            if sys_overview.get('architecture'):
                arch_para = doc.add_paragraph()
                arch_para.add_run('Architecture: ').bold = True
                arch_para.add_run(sys_overview['architecture'])
                self._apply_body_style(arch_para)
            
            if sys_overview.get('components'):
                comp_heading = doc.add_heading('Components', level=2)
                self._apply_heading_style(comp_heading, 2)
                for comp in sys_overview['components']:
                    bullet_para = doc.add_paragraph(f'• {comp}')
                    self._apply_body_style(bullet_para)
                    bullet_para.paragraph_format.left_indent = Inches(0.5)
            
            # Add section images
            self._add_section_images(doc, 'system-overview', section_images)
            self._add_section_separator(doc)
        
        # Functional Requirements
        if frd_data.get('functional_requirements'):
            req_heading = doc.add_heading('FUNCTIONAL REQUIREMENTS', level=1)
            self._apply_heading_style(req_heading, 1)
            
            # Requirements table
            req_table = doc.add_table(rows=len(frd_data['functional_requirements']) + 1, cols=5)
            req_table.style = 'Table Grid'
            
            # Header row
            headers = ['Req ID', 'Title', 'Description', 'Priority', 'Acceptance Criteria']
            for i, header in enumerate(headers):
                cell = req_table.rows[0].cells[i]
                cell.text = header
                cell.paragraphs[0].runs[0].font.bold = True
                cell.paragraphs[0].runs[0].font.name = 'Calibri'
                cell.paragraphs[0].runs[0].font.size = Pt(10)
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                self._add_cell_shading(cell, 'E8E8E8')
            
            # Data rows
            for i, req in enumerate(frd_data['functional_requirements'], 1):
                row = req_table.rows[i]
                row.cells[0].text = req.get('req_id', '')
                row.cells[1].text = req.get('title', '')
                row.cells[2].text = req.get('description', '')
                row.cells[3].text = req.get('priority', '')
                row.cells[4].text = req.get('acceptance_criteria', '')
                
                for cell in row.cells:
                    cell.paragraphs[0].runs[0].font.name = 'Calibri'
                    cell.paragraphs[0].runs[0].font.size = Pt(9)
            
            # Add section images
            self._add_section_images(doc, 'functional-requirements', section_images)
            self._add_section_separator(doc)
        
        # Data Requirements
        if frd_data.get('data_requirements'):
            data_heading = doc.add_heading('DATA REQUIREMENTS', level=1)
            self._apply_heading_style(data_heading, 1)
            
            data_req = frd_data['data_requirements']
            for key, value in data_req.items():
                data_para = doc.add_paragraph()
                data_para.add_run(f'{key.replace("_", " ").title()}: ').bold = True
                data_para.add_run(str(value))
                self._apply_body_style(data_para)
            
            # Add section images
            self._add_section_images(doc, 'data-requirements', section_images)
            self._add_section_separator(doc)
        
        # Interface Requirements
        if frd_data.get('interface_requirements'):
            int_heading = doc.add_heading('INTERFACE REQUIREMENTS', level=1)
            self._apply_heading_style(int_heading, 1)
            
            int_req = frd_data['interface_requirements']
            for key, value in int_req.items():
                int_para = doc.add_paragraph()
                int_para.add_run(f'{key.replace("_", " ").title()}: ').bold = True
                int_para.add_run(str(value))
                self._apply_body_style(int_para)
            
            # Add section images
            self._add_section_images(doc, 'interface-requirements', section_images)
            self._add_section_separator(doc)
        
        # Integration Requirements
        if frd_data.get('integration_requirements'):
            integ_heading = doc.add_heading('INTEGRATION REQUIREMENTS', level=1)
            self._apply_heading_style(integ_heading, 1)
            
            # Integration table
            integ_table = doc.add_table(rows=len(frd_data['integration_requirements']) + 1, cols=4)
            integ_table.style = 'Table Grid'
            
            # Header row
            headers = ['System', 'Method', 'Data Format', 'Frequency']
            for i, header in enumerate(headers):
                cell = integ_table.rows[0].cells[i]
                cell.text = header
                cell.paragraphs[0].runs[0].font.bold = True
                cell.paragraphs[0].runs[0].font.name = 'Calibri'
                cell.paragraphs[0].runs[0].font.size = Pt(11)
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                self._add_cell_shading(cell, 'E8E8E8')
            
            # Data rows
            for i, integ in enumerate(frd_data['integration_requirements'], 1):
                row = integ_table.rows[i]
                row.cells[0].text = integ.get('system', '')
                row.cells[1].text = integ.get('method', '')
                row.cells[2].text = integ.get('data_format', '')
                row.cells[3].text = integ.get('frequency', '')
                
                for cell in row.cells:
                    cell.paragraphs[0].runs[0].font.name = 'Calibri'
                    cell.paragraphs[0].runs[0].font.size = Pt(10)
            
            # Add section images
            self._add_section_images(doc, 'integration-requirements', section_images)
            self._add_section_separator(doc)
            perf_heading = doc.add_heading('PERFORMANCE REQUIREMENTS', level=1)
            self._apply_heading_style(perf_heading, 1)
            
            perf_req = frd_data['performance_requirements']
            for key, value in perf_req.items():
                perf_para = doc.add_paragraph()
                perf_para.add_run(f'{key.replace("_", " ").title()}: ').bold = True
                perf_para.add_run(str(value))
                self._apply_body_style(perf_para)
            
            # Add section images
            self._add_section_images(doc, 'performance-requirements', section_images)
            self._add_section_separator(doc)
        
        # Security Requirements
        if frd_data.get('security_requirements'):
            sec_heading = doc.add_heading('SECURITY REQUIREMENTS', level=1)
            self._apply_heading_style(sec_heading, 1)
            
            for req in frd_data['security_requirements']:
                bullet_para = doc.add_paragraph(f'• {req}')
                self._apply_body_style(bullet_para)
                bullet_para.paragraph_format.left_indent = Inches(0.5)
            
        # Validation Rules
        if frd_data.get('validation_rules'):
            val_heading = doc.add_heading('VALIDATION RULES', level=1)
            self._apply_heading_style(val_heading, 1)
            
            # Validation table
            val_table = doc.add_table(rows=len(frd_data['validation_rules']) + 1, cols=3)
            val_table.style = 'Table Grid'
            
            # Header row
            headers = ['Field', 'Rule', 'Error Message']
            for i, header in enumerate(headers):
                cell = val_table.rows[0].cells[i]
                cell.text = header
                cell.paragraphs[0].runs[0].font.bold = True
                cell.paragraphs[0].runs[0].font.name = 'Calibri'
                cell.paragraphs[0].runs[0].font.size = Pt(11)
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                self._add_cell_shading(cell, 'E8E8E8')
            
            # Data rows
            for i, val in enumerate(frd_data['validation_rules'], 1):
                row = val_table.rows[i]
                row.cells[0].text = val.get('field', '')
                row.cells[1].text = val.get('rule', '')
                row.cells[2].text = val.get('error_message', '')
                
                for cell in row.cells:
                    cell.paragraphs[0].runs[0].font.name = 'Calibri'
                    cell.paragraphs[0].runs[0].font.size = Pt(10)
            
            # Add section images
            self._add_section_images(doc, 'validation-rules', section_images)
            self._add_section_separator(doc)
        
        # Error Handling
        if frd_data.get('error_handling'):
            err_heading = doc.add_heading('ERROR HANDLING', level=1)
            self._apply_heading_style(err_heading, 1)
            
            # Error handling table
            err_table = doc.add_table(rows=len(frd_data['error_handling']) + 1, cols=4)
            err_table.style = 'Table Grid'
            
            # Header row
            headers = ['Error Type', 'Handling Strategy', 'User Message', 'Logging']
            for i, header in enumerate(headers):
                cell = err_table.rows[0].cells[i]
                cell.text = header
                cell.paragraphs[0].runs[0].font.bold = True
                cell.paragraphs[0].runs[0].font.name = 'Calibri'
                cell.paragraphs[0].runs[0].font.size = Pt(11)
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                self._add_cell_shading(cell, 'E8E8E8')
            
            # Data rows
            for i, err in enumerate(frd_data['error_handling'], 1):
                row = err_table.rows[i]
                row.cells[0].text = err.get('error_type', '')
                row.cells[1].text = err.get('handling_strategy', '')
                row.cells[2].text = err.get('user_message', '')
                row.cells[3].text = err.get('logging', '')
                
                for cell in row.cells:
                    cell.paragraphs[0].runs[0].font.name = 'Calibri'
                    cell.paragraphs[0].runs[0].font.size = Pt(10)
            
            # Add section images
            self._add_section_images(doc, 'error-handling', section_images)
            self._add_section_separator(doc)
        
        # Testing Requirements
        if frd_data.get('testing_requirements'):
            test_heading = doc.add_heading('TESTING REQUIREMENTS', level=1)
            self._apply_heading_style(test_heading, 1)
            
            test_req = frd_data['testing_requirements']
            for key, value in test_req.items():
                test_para = doc.add_paragraph()
                test_para.add_run(f'{key.replace("_", " ").title()}: ').bold = True
                test_para.add_run(str(value))
                self._apply_body_style(test_para)
            
            # Add section images
            self._add_section_images(doc, 'testing-requirements', section_images)
            self._add_section_separator(doc)
        
        # Deployment Requirements
        if frd_data.get('deployment_requirements'):
            deploy_heading = doc.add_heading('DEPLOYMENT REQUIREMENTS', level=1)
            self._apply_heading_style(deploy_heading, 1)
            
            deploy_req = frd_data['deployment_requirements']
            for key, value in deploy_req.items():
                deploy_para = doc.add_paragraph()
                deploy_para.add_run(f'{key.replace("_", " ").title()}: ').bold = True
                deploy_para.add_run(str(value))
                self._apply_body_style(deploy_para)
            
            # Add section images
            self._add_section_images(doc, 'deployment-requirements', section_images)
            self._add_section_separator(doc)
    
    def _export_frd_pdf_enhanced(self, frd_data, coverage_data, section_images=None):
        """Export FRD PDF with enhanced styling"""
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"Functional_Requirements_Document_{timestamp}.pdf"
        filepath = os.path.join(self.temp_dir, filename)
        
        doc = SimpleDocTemplate(filepath, pagesize=letter)
        styles = getSampleStyleSheet()
        story = []
        
        # Custom styles
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=20,
            textColor=colors.HexColor('#1e40af'),
            spaceAfter=30,
            alignment=1,
            fontName='Helvetica-Bold'
        )
        
        # Cover page
        if os.path.exists(self.logo_path):
            logo = RLImage(self.logo_path, width=3*inch, height=1.2*inch)
            story.append(logo)
        
        story.append(Spacer(1, 50))
        story.append(Paragraph("FUNCTIONAL REQUIREMENTS DOCUMENT", title_style))
        story.append(Spacer(1, 20))
        
        if frd_data.get('system_overview', {}).get('architecture'):
            story.append(Paragraph(frd_data['system_overview']['architecture'][:150], styles['Normal']))
        
        story.append(PageBreak())
        
        # Dashboard
        if coverage_data:
            self._add_frd_pdf_corporate_dashboard(story, coverage_data, styles)
            story.append(PageBreak())
        
        # Content
        self._add_frd_pdf_corporate_content(story, frd_data, styles)
        
        doc.build(story)
        return filepath
    
    def _add_frd_pdf_corporate_dashboard(self, story, coverage_data, styles):
        """FRD PDF corporate dashboard"""
        story.append(Paragraph("FRD COVERAGE ANALYSIS", styles['Heading1']))
        story.append(Spacer(1, 12))
        
        # System complexity
        if coverage_data.get('system_complexity'):
            story.append(Paragraph(f"<b>System Complexity:</b> {coverage_data.get('system_complexity', 'Medium')}", styles['Normal']))
            story.append(Spacer(1, 20))
        
        # Coverage metrics
        present = len(coverage_data.get('coverage_analysis', {}).get('present_elements', []))
        missing = len(coverage_data.get('coverage_analysis', {}).get('missing_elements', []))
        
        metrics_data = [
            ['Metric', 'Value'],
            ['Coverage Score', f"{coverage_data.get('overall_score', 0)}%"],
            ['Covered Elements', str(present)],
            ['Missing Elements', str(missing)],
            ['Enterprise Readiness', coverage_data.get('enterprise_readiness', 'Unknown')]
        ]
        
        metrics_table = Table(metrics_data, colWidths=[3*inch, 2*inch])
        metrics_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#E8E8E8')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.black),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, -1), 11),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 12),
            ('GRID', (0, 0), (-1, -1), 0.75, colors.black)
        ]))
        
        story.append(metrics_table)
        story.append(Spacer(1, 20))
    
    def _add_frd_pdf_corporate_content(self, story, frd_data, styles):
        """FRD PDF corporate content"""
        # System Overview
        if frd_data.get('system_overview'):
            story.append(Paragraph("SYSTEM OVERVIEW", styles['Heading1']))
            story.append(Spacer(1, 12))
            
            sys_overview = frd_data['system_overview']
            if sys_overview.get('architecture'):
                story.append(Paragraph(f"<b>Architecture:</b> {sys_overview['architecture']}", styles['Normal']))
            
            story.append(Spacer(1, 20))
        
        # Functional Requirements
        if frd_data.get('functional_requirements'):
            story.append(Paragraph("FUNCTIONAL REQUIREMENTS", styles['Heading1']))
            story.append(Spacer(1, 6))
            
            for req in frd_data['functional_requirements']:
                story.append(Paragraph(f"<b>{req.get('req_id', '')}:</b> {req.get('title', '')}", styles['Normal']))
                story.append(Paragraph(req.get('description', ''), styles['Normal']))
            
            story.append(Spacer(1, 16))
    def _export_srd_word_corporate(self, srd_data, coverage_data, section_images=None):
        """Export SRD Word with corporate formatting"""
        doc = Document()
        
        # Set document margins and page setup
        self._setup_document_formatting(doc)
        
        # Add page borders
        self._add_page_borders(doc)
        
        # Add header and footer for SRD
        self._add_srd_header_footer(doc)
        
        # Cover Page
        self._add_srd_corporate_cover_page(doc, srd_data)
        doc.add_page_break()
        
        # Document Control
        self._add_srd_document_control(doc)
        doc.add_page_break()
        
        # Dashboard
        if coverage_data:
            self._add_srd_corporate_dashboard(doc, coverage_data)
            doc.add_page_break()
        
        # SRD Content with corporate styling
        self._add_srd_corporate_content(doc, srd_data, section_images)
        
        # Save
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"System_Requirements_Document_{timestamp}.docx"
        filepath = os.path.join(self.temp_dir, filename)
        doc.save(filepath)
        
        return filepath
    
    def _add_srd_header_footer(self, doc):
        """Add SRD corporate header and footer"""
        try:
            section = doc.sections[0]
            
            # Header
            header = section.header
            header_para = header.paragraphs[0]
            header_para.clear()
            
            # Left: Company name
            left_run = header_para.add_run('ANAND RATHI IT PVT. LTD.')
            left_run.font.name = 'Calibri'
            left_run.font.size = Pt(10)
            left_run.font.bold = True
            left_run.font.color.rgb = RGBColor(*self.colors['primary'])
            
            # Center: Document title
            header_para.add_run('\t\tSystem Requirements Document')
            center_run = header_para.runs[-1]
            center_run.font.name = 'Calibri'
            center_run.font.size = Pt(10)
            center_run.font.bold = True
            
            # Right: Version
            header_para.add_run('\t\tVersion: 1.0')
            right_run = header_para.runs[-1]
            right_run.font.name = 'Calibri'
            right_run.font.size = Pt(10)
            
            # Add bottom border to header
            self._add_header_footer_border(header_para, 'bottom')
            
            # Footer
            footer = section.footer
            footer_para = footer.paragraphs[0]
            footer_para.clear()
            footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            footer_run = footer_para.add_run('Confidential – For Internal Use Only | Page ')
            footer_run.font.name = 'Calibri'
            footer_run.font.size = Pt(9)
            footer_run.font.color.rgb = RGBColor(*self.colors['secondary'])
            
            # Add top border to footer
            self._add_header_footer_border(footer_para, 'top')
            
        except Exception:
            pass  # Skip if SRD header/footer fails
    
    def _add_srd_corporate_cover_page(self, doc, srd_data):
        """Add SRD corporate cover page"""
        # Logo
        if os.path.exists(self.logo_path):
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para.add_run()
            run.add_picture(self.logo_path, width=Inches(3.0))
        
        # Spacing
        for _ in range(3):
            doc.add_paragraph()
        
        # Title
        title = doc.add_heading('SYSTEM REQUIREMENTS DOCUMENT', level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title.runs[0]
        title_run.font.name = 'Calibri'
        title_run.font.size = Pt(20)
        title_run.font.color.rgb = RGBColor(*self.colors['primary'])
        title_run.bold = True
        
        # System Architecture
        if srd_data.get('system_architecture', {}).get('overview'):
            subtitle = doc.add_paragraph(srd_data['system_architecture']['overview'][:80])
            subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
            subtitle_run = subtitle.runs[0]
            subtitle_run.font.name = 'Calibri'
            subtitle_run.font.size = Pt(14)
            subtitle_run.font.color.rgb = RGBColor(*self.colors['secondary'])
            subtitle_run.italic = True
        
        # Spacing
        for _ in range(4):
            doc.add_paragraph()
        
        # Version History Table
        self._add_version_history_table(doc)
    
    def _add_srd_document_control(self, doc):
        """Add SRD document control section"""
        control_heading = doc.add_heading('DOCUMENT CONTROL', level=1)
        self._apply_heading_style(control_heading, 1)
        
        # Document Info Table
        info_table = doc.add_table(rows=6, cols=2)
        info_table.style = 'Table Grid'
        
        info_data = [
            ('Document Title', 'System Requirements Document'),
            ('Document ID', 'SRD-001'),
            ('Version', '1.0'),
            ('Author', 'System Architect'),
            ('Department', 'System Architecture'),
            ('Date', datetime.now().strftime('%Y-%m-%d'))
        ]
        
        for i, (label, value) in enumerate(info_data):
            row = info_table.rows[i]
            row.cells[0].text = label
            row.cells[1].text = value
            
            # Style cells
            row.cells[0].paragraphs[0].runs[0].font.bold = True
            row.cells[0].paragraphs[0].runs[0].font.name = 'Calibri'
            row.cells[0].paragraphs[0].runs[0].font.size = Pt(11)
            row.cells[1].paragraphs[0].runs[0].font.name = 'Calibri'
            row.cells[1].paragraphs[0].runs[0].font.size = Pt(11)
            
            # Add light gray background to header column
            self._add_cell_shading(row.cells[0], 'E8E8E8')
    
    def _add_srd_corporate_dashboard(self, doc, coverage_data):
        """Add SRD corporate dashboard"""
        # Coverage metrics heading
        metrics_heading = doc.add_heading('SRD COVERAGE ANALYSIS', level=1)
        self._apply_heading_style(metrics_heading, 1)
        
        # System complexity
        if coverage_data.get('system_complexity'):
            complexity_para = doc.add_paragraph()
            complexity_para.add_run('System Complexity: ').bold = True
            complexity_para.add_run(coverage_data.get('system_complexity', 'Medium'))
            self._apply_body_style(complexity_para)
        
        # Coverage metrics table
        metrics_table = doc.add_table(rows=4, cols=2)
        metrics_table.style = 'Table Grid'
        
        present = len(coverage_data.get('coverage_analysis', {}).get('present_elements', []))
        missing = len(coverage_data.get('coverage_analysis', {}).get('missing_elements', []))
        
        metrics_data = [
            ('Coverage Score', f"{coverage_data.get('overall_score', 0)}%"),
            ('Covered Elements', str(present)),
            ('Missing Elements', str(missing)),
            ('Enterprise Readiness', coverage_data.get('enterprise_readiness', 'Unknown'))
        ]
        
        for i, (label, value) in enumerate(metrics_data):
            row = metrics_table.rows[i]
            row.cells[0].text = label
            row.cells[1].text = value
            
            # Style cells
            for cell in row.cells:
                cell.paragraphs[0].runs[0].font.name = 'Calibri'
                cell.paragraphs[0].runs[0].font.size = Pt(11)
            
            row.cells[0].paragraphs[0].runs[0].font.bold = True
            self._add_cell_shading(row.cells[0], 'E8E8E8')
    
    def _add_srd_corporate_content(self, doc, srd_data, section_images=None):
        """Add SRD corporate content sections"""
        # System Architecture
        if srd_data.get('system_architecture'):
            arch_heading = doc.add_heading('SYSTEM ARCHITECTURE', level=1)
            self._apply_heading_style(arch_heading, 1)
            
            arch = srd_data['system_architecture']
            if arch.get('overview'):
                overview_para = doc.add_paragraph()
                overview_para.add_run('Overview: ').bold = True
                overview_para.add_run(arch['overview'])
                self._apply_body_style(overview_para)
            
            if arch.get('components'):
                comp_heading = doc.add_heading('Components', level=2)
                self._apply_heading_style(comp_heading, 2)
                for comp in arch['components']:
                    bullet_para = doc.add_paragraph(f'• {comp}')
                    self._apply_body_style(bullet_para)
                    bullet_para.paragraph_format.left_indent = Inches(0.5)
            
            # Add section images
            self._add_section_images(doc, 'system-architecture', section_images)
            self._add_section_separator(doc)
        
        # Hardware Requirements
        if srd_data.get('hardware_requirements'):
            hw_heading = doc.add_heading('HARDWARE REQUIREMENTS', level=1)
            self._apply_heading_style(hw_heading, 1)
            
            hw_req = srd_data['hardware_requirements']
            for key, value in hw_req.items():
                hw_para = doc.add_paragraph()
                hw_para.add_run(f'{key.replace("_", " ").title()}: ').bold = True
                hw_para.add_run(str(value))
                self._apply_body_style(hw_para)
            
        # Software Requirements
        if srd_data.get('software_requirements'):
            sw_heading = doc.add_heading('SOFTWARE REQUIREMENTS', level=1)
            self._apply_heading_style(sw_heading, 1)
            
            sw_req = srd_data['software_requirements']
            for key, value in sw_req.items():
                sw_para = doc.add_paragraph()
                sw_para.add_run(f'{key.replace("_", " ").title()}: ').bold = True
                sw_para.add_run(str(value))
                self._apply_body_style(sw_para)
            
            # Add section images
            self._add_section_images(doc, 'software-requirements', section_images)
            self._add_section_separator(doc)
        
        # Network Requirements
        if srd_data.get('network_requirements'):
            net_heading = doc.add_heading('NETWORK REQUIREMENTS', level=1)
            self._apply_heading_style(net_heading, 1)
            
            net_req = srd_data['network_requirements']
            for key, value in net_req.items():
                net_para = doc.add_paragraph()
                net_para.add_run(f'{key.replace("_", " ").title()}: ').bold = True
                net_para.add_run(str(value))
                self._apply_body_style(net_para)
            
            # Add section images
            self._add_section_images(doc, 'network-requirements', section_images)
            self._add_section_separator(doc)
        
        # Database Requirements
        if srd_data.get('database_requirements'):
            db_heading = doc.add_heading('DATABASE REQUIREMENTS', level=1)
            self._apply_heading_style(db_heading, 1)
            
            db_req = srd_data['database_requirements']
            for key, value in db_req.items():
                db_para = doc.add_paragraph()
                db_para.add_run(f'{key.replace("_", " ").title()}: ').bold = True
                db_para.add_run(str(value))
                self._apply_body_style(db_para)
            
            # Add section images
            self._add_section_images(doc, 'database-requirements', section_images)
            self._add_section_separator(doc)
        
        # Performance Specifications
        if srd_data.get('performance_specifications'):
            perf_heading = doc.add_heading('PERFORMANCE SPECIFICATIONS', level=1)
            self._apply_heading_style(perf_heading, 1)
            
            perf_spec = srd_data['performance_specifications']
            for key, value in perf_spec.items():
                perf_para = doc.add_paragraph()
                perf_para.add_run(f'{key.replace("_", " ").title()}: ').bold = True
                perf_para.add_run(str(value))
                self._apply_body_style(perf_para)
            
            # Add section images
            self._add_section_images(doc, 'performance-specifications', section_images)
            self._add_section_separator(doc)
        
        # System Interfaces
        if srd_data.get('system_interfaces'):
            int_heading = doc.add_heading('SYSTEM INTERFACES', level=1)
            self._apply_heading_style(int_heading, 1)
            
            # Interfaces table
            int_table = doc.add_table(rows=len(srd_data['system_interfaces']) + 1, cols=4)
            int_table.style = 'Table Grid'
            
            # Header row
            headers = ['Interface', 'Type', 'Protocol', 'Data Format']
            for i, header in enumerate(headers):
                cell = int_table.rows[0].cells[i]
                cell.text = header
                cell.paragraphs[0].runs[0].font.bold = True
                cell.paragraphs[0].runs[0].font.name = 'Calibri'
                cell.paragraphs[0].runs[0].font.size = Pt(11)
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                self._add_cell_shading(cell, 'E8E8E8')
            
            # Data rows
            for i, interface in enumerate(srd_data['system_interfaces'], 1):
                row = int_table.rows[i]
                row.cells[0].text = interface.get('interface', '')
                row.cells[1].text = interface.get('type', '')
                row.cells[2].text = interface.get('protocol', '')
                row.cells[3].text = interface.get('data_format', '')
                
                for cell in row.cells:
                    cell.paragraphs[0].runs[0].font.name = 'Calibri'
                    cell.paragraphs[0].runs[0].font.size = Pt(10)
            
        # Security Architecture
        if srd_data.get('security_architecture'):
            sec_heading = doc.add_heading('SECURITY ARCHITECTURE', level=1)
            self._apply_heading_style(sec_heading, 1)
            
            sec_arch = srd_data['security_architecture']
            for key, value in sec_arch.items():
                sec_para = doc.add_paragraph()
                sec_para.add_run(f'{key.replace("_", " ").title()}: ').bold = True
                sec_para.add_run(str(value))
                self._apply_body_style(sec_para)
            
            # Add section images
            self._add_section_images(doc, 'security-architecture', section_images)
            self._add_section_separator(doc)
        
        # Backup & Recovery
        if srd_data.get('backup_recovery'):
            backup_heading = doc.add_heading('BACKUP & RECOVERY', level=1)
            self._apply_heading_style(backup_heading, 1)
            
            backup_req = srd_data['backup_recovery']
            for key, value in backup_req.items():
                backup_para = doc.add_paragraph()
                backup_para.add_run(f'{key.replace("_", " ").title()}: ').bold = True
                backup_para.add_run(str(value))
                self._apply_body_style(backup_para)
            
            # Add section images
            self._add_section_images(doc, 'backup-recovery', section_images)
            self._add_section_separator(doc)
        
        # Monitoring & Logging
        if srd_data.get('monitoring_logging'):
            mon_heading = doc.add_heading('MONITORING & LOGGING', level=1)
            self._apply_heading_style(mon_heading, 1)
            
            mon_req = srd_data['monitoring_logging']
            for key, value in mon_req.items():
                mon_para = doc.add_paragraph()
                mon_para.add_run(f'{key.replace("_", " ").title()}: ').bold = True
                mon_para.add_run(str(value))
                self._apply_body_style(mon_para)
            
            # Add section images
            self._add_section_images(doc, 'monitoring-logging', section_images)
            self._add_section_separator(doc)
        
        # Scalability Requirements
        if srd_data.get('scalability_requirements'):
            scale_heading = doc.add_heading('SCALABILITY REQUIREMENTS', level=1)
            self._apply_heading_style(scale_heading, 1)
            
            scale_req = srd_data['scalability_requirements']
            for key, value in scale_req.items():
                scale_para = doc.add_paragraph()
                scale_para.add_run(f'{key.replace("_", " ").title()}: ').bold = True
                scale_para.add_run(str(value))
                self._apply_body_style(scale_para)
            
            # Add section images
            self._add_section_images(doc, 'scalability-requirements', section_images)
            self._add_section_separator(doc)
        
        # Compliance Standards
        if srd_data.get('compliance_standards'):
            comp_heading = doc.add_heading('COMPLIANCE & STANDARDS', level=1)
            self._apply_heading_style(comp_heading, 1)
            
            for standard in srd_data['compliance_standards']:
                comp_para = doc.add_paragraph()
                comp_para.add_run(f"{standard.get('standard', '')}: ").bold = True
                comp_para.add_run(standard.get('description', ''))
                self._apply_body_style(comp_para)
            
            # Add section images
            self._add_section_images(doc, 'compliance-standards', section_images)
            self._add_section_separator(doc)
    
    def _export_srd_pdf_enhanced(self, srd_data, coverage_data, section_images=None):
        """Export SRD PDF with enhanced styling"""
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"System_Requirements_Document_{timestamp}.pdf"
        filepath = os.path.join(self.temp_dir, filename)
        
        doc = SimpleDocTemplate(filepath, pagesize=letter)
        styles = getSampleStyleSheet()
        story = []
        
        # Custom styles
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=20,
            textColor=colors.HexColor('#1e40af'),
            spaceAfter=30,
            alignment=1,
            fontName='Helvetica-Bold'
        )
        
        # Cover page
        if os.path.exists(self.logo_path):
            logo = RLImage(self.logo_path, width=3*inch, height=1.2*inch)
            story.append(logo)
        
        story.append(Spacer(1, 50))
        story.append(Paragraph("SYSTEM REQUIREMENTS DOCUMENT", title_style))
        story.append(Spacer(1, 20))
        
        if srd_data.get('system_architecture', {}).get('overview'):
            story.append(Paragraph(srd_data['system_architecture']['overview'][:150], styles['Normal']))
        
        story.append(PageBreak())
        
        # Dashboard
        if coverage_data:
            self._add_srd_pdf_corporate_dashboard(story, coverage_data, styles)
            story.append(PageBreak())
        
        # Content
        self._add_srd_pdf_corporate_content(story, srd_data, styles)
        
        doc.build(story)
        return filepath
    
    def _add_srd_pdf_corporate_dashboard(self, story, coverage_data, styles):
        """SRD PDF corporate dashboard"""
        story.append(Paragraph("SRD COVERAGE ANALYSIS", styles['Heading1']))
        story.append(Spacer(1, 12))
        
        # System complexity
        if coverage_data.get('system_complexity'):
            story.append(Paragraph(f"<b>System Complexity:</b> {coverage_data.get('system_complexity', 'Medium')}", styles['Normal']))
            story.append(Spacer(1, 20))
        
        # Coverage metrics
        present = len(coverage_data.get('coverage_analysis', {}).get('present_elements', []))
        missing = len(coverage_data.get('coverage_analysis', {}).get('missing_elements', []))
        
        metrics_data = [
            ['Metric', 'Value'],
            ['Coverage Score', f"{coverage_data.get('overall_score', 0)}%"],
            ['Covered Elements', str(present)],
            ['Missing Elements', str(missing)],
            ['Enterprise Readiness', coverage_data.get('enterprise_readiness', 'Unknown')]
        ]
        
        metrics_table = Table(metrics_data, colWidths=[3*inch, 2*inch])
        metrics_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#E8E8E8')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.black),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, -1), 11),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 12),
            ('GRID', (0, 0), (-1, -1), 0.75, colors.black)
        ]))
        
        story.append(metrics_table)
        story.append(Spacer(1, 20))
    
    def _add_srd_pdf_corporate_content(self, story, srd_data, styles):
        """SRD PDF corporate content"""
        # System Architecture
        if srd_data.get('system_architecture'):
            story.append(Paragraph("SYSTEM ARCHITECTURE", styles['Heading1']))
            story.append(Spacer(1, 12))
            
            arch = srd_data['system_architecture']
            if arch.get('overview'):
                story.append(Paragraph(f"<b>Overview:</b> {arch['overview']}", styles['Normal']))
            
            story.append(Spacer(1, 20))
        
        # Hardware Requirements
        if srd_data.get('hardware_requirements'):
            story.append(Paragraph("HARDWARE REQUIREMENTS", styles['Heading1']))
            story.append(Spacer(1, 6))
            
            hw_req = srd_data['hardware_requirements']
            for key, value in hw_req.items():
                story.append(Paragraph(f"<b>{key.replace('_', ' ').title()}:</b> {value}", styles['Normal']))
            
            story.append(Spacer(1, 16))
    
    def _add_section_images(self, doc, section_id, section_images):
        """Add images for a specific section to the Word document"""
        print(f"DEBUG: _add_section_images called with section_id: {section_id}")
        print(f"DEBUG: section_images keys: {list(section_images.keys()) if section_images else 'None'}")
        
        if not section_images or section_id not in section_images:
            print(f"DEBUG: No images found for section {section_id}")
            return
        
        images = section_images[section_id]
        if not images:
            print(f"DEBUG: Empty images list for section {section_id}")
            return
        
        print(f"DEBUG: Found {len(images)} images for section {section_id}")
        
        for i, image_data in enumerate(images):
            try:
                print(f"DEBUG: Processing image {i+1} for section {section_id}")
                
                # Decode base64 image
                base64_string = image_data.get('data', '')
                if base64_string.startswith('data:image'):
                    # Remove data URL prefix
                    base64_string = base64_string.split(',')[1]
                
                # Convert base64 to bytes
                image_bytes = base64.b64decode(base64_string)
                image_stream = io.BytesIO(image_bytes)
                
                # Add image to document
                image_para = doc.add_paragraph()
                run = image_para.add_run()
                run.add_picture(image_stream, width=Inches(4))
                image_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                print(f"DEBUG: Successfully added image {i+1} to section {section_id}")
                
                # Add caption if available
                caption = image_data.get('caption', '')
                if caption:
                    caption_para = doc.add_paragraph(caption)
                    caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    caption_run = caption_para.runs[0]
                    caption_run.font.name = 'Calibri'
                    caption_run.font.size = Pt(10)
                    caption_run.font.italic = True
                    caption_run.font.color.rgb = RGBColor(*self.colors['secondary'])
                
                # Add spacing after image
                doc.add_paragraph()
                
            except Exception as e:
                # Skip problematic images
                print(f"ERROR: Failed to add image {i+1} to section {section_id}: {str(e)}")
                continue